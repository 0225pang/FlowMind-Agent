"""
生成《系统设计文档》Word 文档
FlowMind Agent：基于飞书生态的 AI 内容运营与知识管理智能体平台
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── 样式定义 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading styles
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '黑体'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        heading_style.font.size = Pt(16)
    elif level == 2:
        heading_style.font.size = Pt(14)
    else:
        heading_style.font.size = Pt(13)

# ═══════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('系统设计文档')
run.font.size = Pt(26)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('FlowMind Agent\n基于飞书生态的 AI 内容运营与知识管理智能体平台')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run('版本：v1.0\n日期：2026年6月\n状态：课程设计 Demo').font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 摘要
# ═══════════════════════════════════════════════
doc.add_heading('摘  要', level=1)

doc.add_paragraph(
    'FlowMind Agent 是一个面向内容运营、教育服务、个人 IP 运营和小型团队协作场景的 AI 智能体平台。'
    '系统以 AI 对话工作台作为统一入口，同时提供 Web 前端、桌面客户端（Python/PySide6）和 Android 移动客户端三套界面，'
    '结合飞书生态中的文档、多维表格、任务协作和群机器人能力，'
    '并结合本地业务数据库、向量数据库和大模型 API，实现内容选题生成、资料库整理、知识问答、学员信息管理、'
    '院校项目情报管理、数据可视化和任务协同。'
)
doc.add_paragraph(
    '本文档为 FlowMind Agent 的《系统设计文档》，面向课程结课设计提交。文档详细阐述了系统的技术选型、'
    '总体架构设计、组件与模块结构（含三套客户端）、数据架构设计、服务协作与事务设计、认证鉴权与安全设计、高并发与性能设计、'
    '日志监控与故障定位、部署与运维设计，以及 12 个技术关键点的详细解决方案。'
)
doc.add_paragraph(
    '系统采用前后端分离架构：三套客户端（Web 基于 Vue3 + TypeScript + Vite + Element Plus；'
    '桌面基于 Python/PySide6；移动基于 Android/Java）通过统一 REST/SSE API 访问同一套后端。'
    '后端基于 Spring Boot 3.x 多模块 Maven 工程。Demo 阶段使用 app-service 聚合启动，同时保留清晰的微服务边界，后续可平滑拆分为独立微服务。'
    '数据存储采用 MySQL 作为业务事实源，Weaviate 作为向量检索引擎，Redis 和 MinIO 预留后续扩展。'
    'AI 智能体层通过 LLMClient 抽象、Agent Router 自动路由和 AgentExtension 工具扩展机制，实现了大模型厂商无关、'
    '工具能力解耦、AI 输出可追溯的设计目标。'
)
doc.add_paragraph(
    '关键词：AI 智能体；飞书生态；内容运营；知识管理；向量检索；微服务架构；SSE 流式输出；'
    '桌面客户端；Android 移动端；PySide6'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 目录占位
# ═══════════════════════════════════════════════
doc.add_heading('目  录', level=1)
doc.add_paragraph('（在 Word 中插入自动目录：引用 → 目录 → 自动目录）')
doc.add_page_break()

# ═══════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════

def add_heading(text, level=1):
    """添加标题并返回"""
    return doc.add_heading(text, level=level)

def add_para(text, bold=False):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    if bold:
        run.font.bold = True
    return p

def add_table(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_code_block(code, language=""):
    """添加代码块（用表格模拟）"""
    p = doc.add_paragraph()
    run = p.add_run(f'[{language}]' if language else '')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

def add_figure(title, caption, code, lang="Mermaid"):
    """添加图：标题 + 代码 + 图注"""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'【{title}】')
    run.font.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if code:
        add_code_block(code, lang)

    p = doc.add_paragraph()
    run = p.add_run(f'图注：{caption}')
    run.font.size = Pt(10)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ═══════════════════════════════════════════════
# 第一章 技术选型
# ═══════════════════════════════════════════════
add_heading('1  技术选型', 1)

# 1.1
add_heading('1.1  客户端技术选型', 2)

add_para(
    'FlowMind Agent 提供三套客户端界面，覆盖 Web 浏览器、桌面操作系统和 Android 移动设备三种使用场景。'
    '三端通过统一的 REST/SSE API 访问同一套后端服务，共享全部业务逻辑和数据。'
)
add_para('Web 前端采用 Vue 3 + TypeScript + Vite + Element Plus 技术栈，构建面向内容运营人员的 AI 工作台与后台管理系统。'
         '桌面客户端采用 Python/PySide6 + QSS 样式系统，高保真复刻 Web 端全部 10 个业务页面，并额外支持离线 Demo 模式和双渲染器兜底（Tkinter）。'
         'Android 移动客户端采用原生 Java 开发（MainActivity + ApiClient），支持 SSE 流式对话和核心业务功能。',
         bold=True)

add_heading('1.1.1  Web 前端已选技术及选型理由', 3)

add_table(
    ['技术', '版本/说明', '选型理由'],
    [
        ['Vue 3', '3.x',
         '渐进式框架，Composition API 便于组织复杂交互逻辑；响应式系统适合 AI 工作台的多状态管理场景；'
         '生态成熟，社区活跃，适合快速构建交互式管理系统。对比 React，Vue 的模板语法对于后台管理系统的'
         '表格、表单、弹窗等场景更直观，学习曲线更平缓。'],
        ['TypeScript', '5.x',
         '静态类型检查提高接口数据结构和组件状态的可靠性；AI 工作台中涉及大量 Agent 消息、SSE 事件、'
         'Tool Trace 等复杂数据结构，TypeScript 能有效减少运行时类型错误。'],
        ['Vite', '5.x',
         '基于 ESM 的构建工具，冷启动快、HMR 极速，适合课程设计阶段的快速迭代开发。相比 Webpack，'
         'Vite 的开发体验显著更优，无需额外配置即可获得 TypeScript、CSS 预处理等支持。'],
        ['Element Plus', '2.x',
         '组件丰富（表格、表单、弹窗、抽屉、标签页、日历、消息提示等），覆盖后台管理系统 90% 以上的交互需求；'
         '与 Vue 3 深度集成，文档完善，可定制主题。对比 Ant Design Vue，Element Plus 在国内 Vue 3 生态中'
         '使用更广泛，示例更丰富。'],
        ['ECharts', '5.x',
         '适合数据分析页面（学员分布、GPA 统计、申请漏斗、内容发布统计等）；图表类型丰富，支持交互式数据探索。'],
        ['Pinia', '2.x',
         'Vue 3 官方推荐状态管理库，轻量且类型安全；用于管理用户登录态、Agent 会话列表、SSE 连接状态等全局状态。'],
        ['Vue Router', '4.x',
         'Vue 3 官方路由库；管理 AI 工作台、内容管理、学员管理、院校情报、数据分析等页面的前端路由。'],
        ['SSE (Server-Sent Events)', '—',
         '用于 AI 流式回答的单向推送协议；比 WebSocket 更轻量，浏览器原生支持，无需额外库；'
         '适合 LLM 文本流式输出场景（服务端单向推送文本，客户端不需要频繁双向通信）。'],
    ]
)

add_heading('1.1.2  Web 前端替代方案对比', 3)

add_table(
    ['替代方案', '优势', '劣势', '本项目不采用的原因'],
    [
        ['React / Next.js',
         '生态庞大、社区资源多',
         'JSX 语法在后台管理场景不如模板直观；Next.js 的 SSR 对 AI 工作台无显著收益',
         '课程设计团队对 Vue 更熟悉；Vue 的模板语法在表格/表单场景效率更高'],
        ['Ant Design Vue',
         '组件库功能完善、设计规范成熟',
         '体积较大、定制化不如 Element Plus 灵活',
         'Element Plus 对表格/标签页/抽屉等场景的原生支持更贴近本项目需求'],
        ['WebSocket',
         '全双工通信、适合实时协作',
         '协议开销大、需要额外的心跳维护',
         'AI 流式回答是单向推送场景，SSE 更轻量；WebSocket 可能引入不必要的连接管理复杂度'],
        ['原生 HTML/CSS/JS',
         '无框架依赖、包体积小',
         '开发效率低、复杂交互难以维护',
         'AI 工作台涉及多面板、多标签、实时流式展示等复杂交互，不适合原生开发'],
    ]
)

add_heading('1.1.3  桌面与移动客户端技术选型', 3)

add_table(
    ['客户端', '技术栈', '选型理由'],
    [
        ['桌面客户端',
         'Python + PySide6 + QSS 样式系统 + httpx；'
         'Tkinter 作为兜底渲染器',
         'PySide6 提供接近原生的现代 GUI 组件，QSS 支持与 Web CSS 类似的设计语言定义能力；'
         'Python 生态丰富，适合快速原型开发；Tkinter 作为标准库无需额外安装，保证在任何 Python 环境都能运行核心功能'],
        ['Android 移动客户端',
         '原生 Android Java + OkHttp/HttpURLConnection；'
         'Gradle 构建；APK 直接安装',
         '原生开发保证启动速度和滑动流畅性；与 Web 端共享同一套后端 API 和 Token 认证；'
         'ApiClient 封装 SSE 流式解析，移动端也能获得流式对话体验'],
    ]
)

# 1.2
add_heading('1.2  后端技术选型', 2)

add_para(
    'FlowMind Agent 后端采用 Spring Boot 3.x + Java 17 + Maven 多模块工程，以标准 REST API 对外暴露服务，'
    '集成 Swagger/OpenAPI 文档、Spring JDBC 轻量数据访问，并预留 Redis、MinIO 扩展点。'
)

add_heading('1.2.1  已选技术及选型理由', 3)

add_table(
    ['技术', '版本/说明', '选型理由'],
    [
        ['Spring Boot', '3.x',
         '业界主流的 Java 后端框架，自动配置、起步依赖、嵌入式服务器，适合构建标准化 REST 服务；'
         'Spring MVC 的注解驱动开发模式与多模块结构天然匹配。'],
        ['Java', '17 LTS',
         '长期支持版本，提供 Record、Sealed Class、Pattern Matching 等现代语言特性；'
         '与企业主流 JDK 版本一致，方便后续生产部署。'],
        ['Maven 多模块', '—',
         '按服务边界拆分 Maven 模块（ai-agent-service、content-service、knowledge-service 等），'
         '每个模块有独立的 pom.xml 和包结构；Demo 阶段通过 app-service 聚合启动，'
         '但源码层面已经体现了清晰的微服务边界。'],
        ['Spring MVC REST API', '—',
         '基于注解的 Controller 开发模式成熟；统一 JSON 响应格式，Swagger 自动生成 API 文档。'],
        ['Spring JDBC / Mapper', '—',
         '轻量级数据访问方式，直接编写 SQL，避免 JPA/Hibernate 的不必要抽象开销；'
         '适合本项目以 SQL 为主的业务场景，且与 MySQL 特性配合紧密。'],
        ['Swagger / OpenAPI', '3.x',
         '自动生成接口文档，方便前端和团队成员查看 API 定义；支持在线调试。'],
        ['MySQL', '9.x',
         '成熟的关系型数据库，适合存储用户、角色、会话、内容、学员、院校等结构化业务数据；'
         '支持事务、外键、索引，生态成熟度高。'],
        ['Weaviate', '—',
         '开源向量数据库，支持语义检索；通过 GraphQL 接口管理向量对象，适合知识库文档的向量化存储和相似度检索。'],
        ['Redis（预留）', '—',
         '计划用于缓存（热点配置、知识库统计）、会话管理、API 限流。Demo 阶段暂不使用，接口预留。'],
        ['MinIO（预留）', '—',
         '计划用于图片（内容配图、学员头像）和文件（知识库附件）的对象存储。Demo 阶段文件存储在本地或 MySQL。'],
        ['OpenAI-Compatible LLM Client', '—',
         '统一封装大模型 API 调用，支持 DeepSeek、OpenAI、豆包、通义千问等任何兼容 OpenAI 协议的服务商；'
         '同时提供 MockLLMClient 用于无 Key 或离线 Demo。'],
        ['lark-cli / 飞书开放平台 API', '—',
         '通过 lark-cli 命令行工具桥接飞书 API；飞书文档/多维表格/任务/机器人等能力统一封装在 feishu-service 和 '
         'LarkCliMcpExtension 中。'],
    ]
)

add_heading('1.2.2  替代方案对比', 3)

add_table(
    ['替代方案', '本项目不采用的原因'],
    [
        ['Node.js / NestJS', 'Java 生态在 Spring Boot 的企业级特性（事务管理、拦截器、多模块）方面更成熟；团队后端成员对 Spring Boot 更熟悉'],
        ['Python / FastAPI', '异步性能好但 Python 的类型系统不如 Java 严谨；Java 在企业级项目中的可维护性和工程化程度更高'],
        ['PostgreSQL', '功能强大但 MySQL 在课程设计环境和 Docker 部署中更轻量、更方便；本项目的数据关系模型不需要 PG 的高级特性'],
        ['MongoDB', 'AI 会话、内容、学员等数据具有明显的结构化特征和关联关系，更适合 MySQL；向量检索已由 Weaviate 承担'],
        ['Elasticsearch', '全文检索能力强大但部署成本高；课程 Demo 规模下，MySQL 索引 + Weaviate 向量检索已满足需求'],
        ['Milvus', '向量检索性能优异但部署资源要求高；Weaviate 更轻量，内置文档管理，适合 Demo 阶段'],
    ]
)

# 1.3
add_heading('1.3  数据库与存储技术选型', 2)

add_para(
    '本系统采用"MySQL 作为业务事实源 + Weaviate 作为向量检索引擎 + Redis/MinIO 预留扩展"的存储架构。'
    '需要特别说明的是：'
)

add_para('① MySQL 并不"过时"：在结构化业务数据、事务、一致性、SQL 查询灵活性、生态成熟度（如 Navicat、DBeaver、数据迁移工具）方面，MySQL 仍然是企业应用的主流选择。本项目的用户、角色、权限、会话、内容选题、学员、院校等数据均具有明确的结构化特征和外键关联，适合使用 MySQL。', bold=False)
add_para('② 向量数据库不是替代 MySQL，而是补充语义检索能力：Weaviate 保存的是知识文档的向量化分片，用于语义相似度搜索。MySQL 仍然保存原始文档元数据、标签、同步日志等。两者通过 mysqlId 字段关联。', bold=False)
add_para('③ Redis 不是数据库主存储：Redis 在本系统中仅作为缓存和高性能中间层使用。会话状态、热点数据、限流计数器等适合放入 Redis，但不替代 MySQL 的业务主数据。', bold=False)
add_para('④ MinIO 不是关系数据库：MinIO 作为 S3 兼容的对象存储，专门处理图片、文件等非结构化二进制数据，与 MySQL 的关系数据互补。', bold=False)

# 1.4
add_heading('1.4  AI 与智能体技术选型', 2)

add_para(
    '本系统的 AI 智能体层设计遵循"厂商无关、工具解耦、输出可追溯"三大原则，具体选型如下：'
)

add_table(
    ['组件', '说明', '设计目标'],
    [
        ['Agent 接口', '定义 execute() 和 stream() 方法，每个智能体实现该接口',
         '统一的智能体执行契约，新增智能体不影响主流程'],
        ['AgentRouter', '根据用户输入自动判断任务类型，路由到对应 Agent',
         '降低用户操作成本——用户不需要手动选择智能体类型'],
        ['LLMClient 接口', '抽象大模型调用，支持 complete() 和 stream()',
         '避免系统强绑定某个大模型厂商'],
        ['MockLLMClient', '返回模拟数据，用于离线 Demo 或无 Key 运行',
         '方便开发调试和课堂演示'],
        ['OpenAiCompatibleLLMClient', '对接任何兼容 OpenAI 协议的大模型 API',
         '支持 DeepSeek、OpenAI、豆包、通义千问等'],
        ['AgentExtension 接口', '定义 name()、type()、description()、runtimeContext()',
         '让工具能力和业务智能体解耦；新增工具只需实现此接口'],
        ['McpToolProvider', '表示外部 MCP 工具能力',
         '标准化工具接入协议'],
        ['SkillProvider', '表示固定工作流或 SOP 能力',
         '封装内容生成 SOP、知识库同步 SOP 等流程'],
        ['AgentTraceService', '记录每次工具调用的输入输出和耗时',
         '让用户知道系统到底调用了哪些工具'],
        ['ConversationEntity (metadata)', '保存 Thinking/Reasoning 过程和工具 Trace',
         'AI 输出可追溯、可审计、可回放'],
    ]
)

# 1.5
add_heading('1.5  技术选型对比与结论', 2)

add_para(
    '综合以上分析，本系统的技术选型以"课程 Demo 快速落地 + 架构可扩展"为核心目标。'
    '前端选择 Vue 3 生态而非 React，主要是考虑到 Vue 的模板语法在后台管理系统的表格/表单/弹窗类场景中开发效率更高，'
    '且团队对 Vue 更熟悉。后端选择 Spring Boot 而非 Node.js 或 Python，主要是因为 Java 在企业级项目中的工程化程度、'
    '事务管理以及多模块拆分能力更适合模拟微服务架构。'
    '数据存储选择 MySQL + Weaviate 的组合而非单一 PostgreSQL 或 MongoDB，是因为结构化业务数据和语义向量检索是两类'
    '不同的存储需求，分开处理更灵活。AI 层通过 LLMClient + AgentRoute + AgentExtension 三层抽象，'
    '保证了对大模型厂商、工具能力和业务场景三个维度的解耦。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第二章 系统总体架构设计
# ═══════════════════════════════════════════════
add_heading('2  系统总体架构设计', 1)

add_heading('2.1  总体架构概述', 2)

add_para(
    'FlowMind Agent 采用经典的前后端分离架构，共提供三套客户端：Web 前端（Vue 3 SPA）、桌面客户端（Python/PySide6）'
    '和 Android 移动客户端（Java）。三端通过统一的 HTTP REST API 和 SSE 协议与同一套后端通信，'
    '共享全部业务逻辑、用户认证和数据存储。'
    '后端是基于 Spring Boot 3.x 的多模块 Maven 工程，Demo 阶段由 app-service 模块聚合启动所有子模块。'
    '逻辑上，后端按照微服务边界划分为 9 个业务服务模块和 3 个公共模块，每个模块有独立的包结构、Controller、Service、Mapper。'
)

add_para(
    '系统目前的支持层次分为四层：'
)
add_para('• 表示层（Presentation）：三套客户端——① Vue 3 Web 前端（Element Plus + ECharts + SSE EventSource）；'
         '② Python/PySide6 桌面客户端（高保真复刻 10 个业务页面，支持离线 Demo）；'
         '③ Android 移动客户端（原生 Java，支持 SSE 流式对话）', bold=False)
add_para('• 网关层（Gateway）：gateway-service 提供统一路由和 CORS 配置', bold=False)
add_para('• 业务服务层（Business Services）：9 个业务模块，每个模块内包含 Controller → Service → Mapper 三层', bold=False)
add_para('• 基础设施层（Infrastructure）：MySQL、Weaviate、LLM API、飞书 Open Platform，以及预留的 Redis 和 MinIO', bold=False)

add_heading('2.2  系统总体架构图', 2)

add_figure(
    '图 2-1  系统总体架构图',
    '该图展示了 FlowMind Agent 系统的完整架构，包括用户层、三套客户端层、网关层、业务服务层、'
    'AI 智能体层和基础设施层。图中箭头表示依赖方向：三套客户端均通过 REST/SSE 访问同一后端，'
    '网关统一路由，业务服务依赖基础设施和 AI 智能体层。',
    '''
┌───────────────────────────────────────────────────────────────────────────┐
│                       用户 (浏览器 / 桌面 / 手机)                           │
└──────────────────────────────────┬────────────────────────────────────────┘
                                   │ HTTPS
              ┌────────────────────┼────────────────────────┐
              ▼                    ▼                        ▼
┌──────────────────────┐┌──────────────────────┐┌──────────────────────┐
│  Web 前端 (Vue 3)    ││ 桌面客户端 (PySide6)  ││ 移动客户端 (Android) │
│  Vite + Element Plus ││  Python + QSS 样式   ││  Java + SSE 流式     │
│  ┌────────┐┌───────┐││  ┌────────┐┌───────┐││  ┌────────┐┌───────┐│
│  │AI工作台││内容管理│││  │AI工作台││内容管理│││  │AI工作台││知识库 ││
│  │SSE流式 ││选题文案│││  │SSE流式 ││选题文案│││  │SSE流式 ││学员   ││
│  └────────┘└───────┘││  └────────┘└───────┘││  └────────┘└───────┘│
└──────────┬───────────┘└──────────┬───────────┘└──────────┬───────────┘
           │                       │                       │
           └───────────────────────┼───────────────────────┘
                                   │ REST / SSE (统一后端 API)
                                   ▼
┌───────────────────────────────────────────────────────────────────────────┐
│                      Nginx / Gateway Service (统一路由 / CORS)              │
└──────────────────────────────────┬────────────────────────────────────────┘
                             │ REST / SSE
                             ▼
┌─────────────────────────────────────────────────────────────────────┐
│                    Gateway Service (统一路由 / CORS)                 │
└────────────────────────────┬────────────────────────────────────────┘
                             │
                             ▼
┌─────────────────────────────────────────────────────────────────────┐
│                   app-service (Demo 聚合启动入口)                    │
│                                                                     │
│  ┌───────────┐ ┌───────────┐ ┌───────────┐ ┌───────────┐          │
│  │user-      │ │ai-agent-  │ │content-   │ │knowledge- │          │
│  │service    │ │service    │ │service    │ │service    │          │
│  │           │ │┌─────────┐│ │┌─────────┐│ │┌─────────┐│          │
│  │• 登录注册 │ ││Agent    ││ ││Port/    ││ ││Embedding││          │
│  │• 角色权限 │ ││Router   ││ ││Adapter  ││ ││Service  ││          │
│  │• RBAC    │ ││LLMClient││ ││SOP 引擎 ││ ││Weaviate ││          │
│  └───────────┘ ││Extension││ │└─────────┘│ ││Client   ││          │
│                │└─────────┘│ └───────────┘ │└─────────┘│          │
│                └───────────┘               └───────────┘          │
│  ┌───────────┐ ┌───────────┐ ┌───────────┐ ┌───────────┐          │
│  │student-   │ │school-    │ │analytics- │ │feishu-    │          │
│  │service    │ │service    │ │service    │ │service    │          │
│  │           │ │           │ │           │ │           │          │
│  │• 画像管理 │ │• 院校情报 │ │• 数据概览 │ │• 飞书文档 │          │
│  │• 申请进度 │ │• 项目匹配 │ │• 图表统计 │ │• 多维表格 │          │
│  └───────────┘ └───────────┘ └───────────┘ │• 群机器人 │          │
│                                            └───────────┘          │
│  ┌───────────┐ ┌───────────┐ ┌───────────┐                        │
│  │common-core│ │common-    │ │common-web │                        │
│  │           │ │security   │ │           │                        │
│  │ApiResponse│ │AuthInter- │ │OpenApi    │                        │
│  │IdGenerator│ │ceptor     │ │Config     │                        │
│  └───────────┘ └───────────┘ └───────────┘                        │
└───────┬──────────┬──────────┬──────────┬──────────┬────────────────┘
        │          │          │          │          │
        ▼          ▼          ▼          ▼          ▼
┌──────────┐┌──────────┐┌──────────┐┌──────────┐┌──────────────────┐
│  MySQL   ││ Weaviate ││  Redis   ││  MinIO   ││   External APIs  │
│  (业务   ││ (向量    ││  (缓存   ││  (对象   ││                  │
│  数据)   ││  检索引擎)││  预留)   ││  存储预留)││ LLM API (兼容    │
│          ││          ││          ││          ││ OpenAI 协议)     │
│          ││          ││          ││          ││ 飞书 Open Platform│
└──────────┘└──────────┘└──────────┘└──────────┘└──────────────────┘
''',
    ''
)

add_heading('2.3  子系统与微服务划分', 2)

add_para(
    'Demo 阶段为何采用 app-service 聚合启动而非真正微服务部署？主要原因有：'
)
add_para('① 降低课程 Demo 的运行复杂度：学生和老师无需启动多个 JVM 进程、配置注册中心和负载均衡即可运行完整系统。', bold=False)
add_para('② 保留微服务边界：虽然共享同一个 JVM，但每个模块有独立的 pom.xml、独立的 Java 包结构，源码层面已经体现了微服务边界。', bold=False)
add_para('③ 平滑拆分路径：当需要拆分为真实微服务时，只需为每个模块编写独立的 Application 启动类和 application.yml，引入 Spring Cloud Gateway / Nacos / Feign 等服务治理组件即可，无需重构业务代码。', bold=False)

add_para(
    '后端共划分 12 个 Maven 模块，其职责和边界如下表所示：'
)

add_heading('2.4  服务边界与数据交换', 2)

add_table(
    ['模块', '职责', '不负责'],
    [
        ['gateway-service', '统一路由入口、CORS 配置、简单鉴权入口；后续可替换为 Spring Cloud Gateway', '不处理具体业务逻辑、不直接访问业务表'],
        ['user-service', '用户登录/注册、用户信息管理、角色权限 CRUD、Demo Token/JWT 签发、RBAC 数据初始化', '不处理内容生成、不处理飞书同步'],
        ['ai-agent-service', 'AI 对话总入口、Agent Router 自动路由、Prompt 模板管理、LLMClient 抽象与实现、AgentExtension 扩展机制、会话历史管理、Thinking 和 Trace 落库', '不直接承担所有业务数据管理、不直接绕过业务服务写业务表'],
        ['knowledge-service', '知识文档 CRUD、标签管理、飞书文档同步记录、文档摘要、Embedding 向量化、Weaviate 向量检索、MySQL fallback', '不直接生成内容运营文案、不管理学院业务'],
        ['content-service', '内容主题库、文案库（含多版本）、内容日历、内容 SOP 引擎、文案评分（1-5 星）、图片/配图管理、内容生成流水、发布效果指标', '不直接管理学院画像、不直接决定飞书底层 API 细节'],
        ['student-service', '学员基本信息、学员画像（GPA/排名/英语/风险等级）、申请进度追踪、AI 分析接口', '不管理内容日历、不管理向量数据库'],
        ['school-service', '院校信息、夏令营/预推免项目、报名条件/材料要求/截止日期、院校推荐匹配', '不管理飞书文档同步、不管理内容文案库'],
        ['analytics-service', '数据概览仪表盘、学员分布统计、GPA 分布、申请漏斗分析、内容发布统计、院校截止趋势', '不直接修改业务主数据（只读分析）'],
        ['feishu-service', '飞书文档读写、飞书多维表格操作、飞书任务协作、飞书群机器人消息、飞书同步日志、lark-cli 桥接', '不决定业务 SOP、不直接替代知识库或内容服务'],
        ['common-core', '统一响应格式 ApiResponse、ID 生成器 IdGenerator、通用工具类', '不包含业务功能'],
        ['common-security', '认证拦截器 AuthInterceptor、Token 工具类 TokenUtil、安全权限服务 SecurityPermissionService', '不包含业务功能'],
        ['common-web', 'Swagger/OpenAPI 文档配置、Web MVC 配置、CORS 全局配置', '不包含业务功能'],
    ]
)

add_heading('2.5  Demo 聚合启动与后续微服务拆分方案', 2)

add_para(
    'Demo 阶段的启动方式：app-service 模块包含主启动类 FlowMindApplication，通过 Maven 依赖引入所有子模块，'
    '执行 mvnw spring-boot:run 即可同时加载所有服务。各模块之间通过 Spring Bean 注入进行内部调用。'
)

add_para(
    '后续拆分为真实微服务的步骤：'
)
add_para('① 为每个业务模块添加独立的 @SpringBootApplication 启动类和 application.yml。', bold=False)
add_para('② 引入 Spring Cloud Gateway 替代 gateway-service 作为统一 API 网关。', bold=False)
add_para('③ 引入 Nacos 或 Consul 作为服务注册与发现中心。', bold=False)
add_para('④ 将模块间的 Spring Bean 调用替换为 Feign 声明式 HTTP 调用或 Spring Cloud LoadBalancer。', bold=False)
add_para('⑤ 跨服务的异步处理（如飞书同步、向量索引更新）可引入 RabbitMQ 或 Kafka 消息队列。', bold=False)
add_para('⑥ 敏感配置从 application-local.yml 迁移到 Nacos Config 或环境变量。', bold=False)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第三章 组件设计与模块结构
# ═══════════════════════════════════════════════
add_heading('3  组件设计与模块结构', 1)

add_heading('3.1  组件图', 2)

add_figure(
    '图 3-1  UML 组件图',
    '该组件图展示了系统的主要组件及其依赖关系。前端组件通过 API Client 访问后端 REST/SSE 接口；'
    '后端 Controller 层调用 Service 层，Service 层通过 Repository/Mapper 访问数据库或通过 Adapter 调用外部系统；'
    'Agent Core 通过 LLMClient 调用大模型，通过 AgentExtension 调用扩展工具。',
    '''
┌───────────────────────────────┐     ┌───────────────────────────────────┐
│      <<Vue Frontend>>        │     │     <<Spring Boot Backend>>        │
│                              │     │                                    │
│ ┌──────────────────────────┐ │     │ ┌──────────────────────────────┐   │
│ │  Vue Router              │ │     │ │ Controllers                 │   │
│ │  (路由守卫 + 页面路由)    │ │     │ │ AgentController             │   │
│ └──────────────────────────┘ │     │ │ ContentController           │   │
│                              │     │ │ KnowledgeController          │   │
│ ┌──────────────────────────┐ │     │ │ StudentController           │   │
│ │  API Client              │─┼──REST/SSE──▶│ SchoolController             │   │
│ │  (fetch + EventSource)   │ │     │ │ AnalyticsController         │   │
│ └──────────────────────────┘ │     │ │ FeishuController            │   │
│                              │     │ │ AuthController              │   │
│ ┌──────────────────────────┐ │     │ └──────────┬───────────────────┘   │
│ │  Agent Workspace         │ │     │            │                       │
│ │  (SSE 流式展示面板)       │ │     │            ▼                       │
│ │  • Thinking 折叠区        │ │     │ ┌──────────────────────────────┐   │
│ │  • Trace 工具调用区       │ │     │ │  Services                    │   │
│ │  • Delta 流式回答区       │ │     │ │  AgentRouter                 │   │
│ └──────────────────────────┘ │     │ │  ConversationService         │   │
│                              │     │ │  ContentLibraryService       │   │
│ ┌──────────────────────────┐ │     │ │  ContentSopService           │   │
│ │  State (Pinia)           │ │     │ │  KnowledgeService            │   │
│ │  auth / sessions / sse   │ │     │ │  ...                         │   │
│ └──────────────────────────┘ │     │ └──────────┬───────────────────┘   │
│                              │     │            │                       │
│ ┌──────────────────────────┐ │     │            ▼                       │
│ │  ECharts Dashboard       │ │     │ ┌──────────────────────────────┐   │
│ │  (数据分析可视化)         │ │     │ │  Mappers / Repositories      │   │
│ └──────────────────────────┘ │     │ │  UserMapper                  │   │
└───────────────────────────────┘     │ │  ContentMapper               │   │
                                      │ │  KnowledgeMapper             │   │
                                      │ │  ConversationMapper          │   │
                                      │ │  SessionMapper               │   │
                                      │ └──────────┬───────────────────┘   │
                                      │            │                       │
                                      │ ┌──────────┴───────────────────┐   │
                                      │ │  Agent Core                  │   │
                                      │ │  ┌────────────────────────┐  │   │
                                      │ │  │ Agent Interface        │  │   │
                                      │ │  │ BaseAgent / ContentAgent│  │   │
                                      │ │  │ KnowledgeAgent / etc.  │  │   │
                                      │ │  └────────────────────────┘  │   │
                                      │ │  ┌────────────────────────┐  │   │
                                      │ │  │ LLMClient              │  │   │
                                      │ │  │ MockLLMClient          │  │   │
                                      │ │  │ OpenAiCompatibleClient │  │   │
                                      │ │  └────────────────────────┘  │   │
                                      │ │  ┌────────────────────────┐  │   │
                                      │ │  │ AgentExtensions        │  │   │
                                      │ │  │ McpToolProvider        │  │   │
                                      │ │  │ SkillProvider          │  │   │
                                      │ │  │ LarkCliMcpExtension    │  │   │
                                      │ │  │ SemanticVectorSearchExt│  │   │
                                      │ │  │ RuntimeToolExtensions  │  │   │
                                      │ │  └────────────────────────┘  │   │
                                      │ └──────────────────────────────┘   │
                                      └──────────┬────────────────────────┘
                                                 │
                          ┌──────────────────────┼──────────────────────┐
                          ▼                      ▼                      ▼
                   ┌──────────┐          ┌──────────┐          ┌──────────────┐
                   │  MySQL   │          │ Weaviate │          │  External API│
                   │  (业务DB)│          │ (向量DB) │          │  LLM / 飞书   │
                   └──────────┘          └──────────┘          └──────────────┘
''',
    ''
)

add_para(
    '组件之间的依赖关系说明：'
)
add_para('① 前端通过 API Client（基于 fetch 和 EventSource）访问后端的 REST 接口和 SSE 流式端点。前端不持有 LLM Key 或飞书 Key，所有外部 API 调用均经过后端代理。', bold=False)
add_para('② Controller 层负责接收 HTTP 请求、参数校验和响应格式化，调用对应的 Service 层方法。Controller 不做业务逻辑判断。', bold=False)
add_para('③ Service 层是业务逻辑的核心：AgentRouter 根据用户输入内容自动路由到对应的 Agent；ContentSopService 执行内容 SOP 流程；KnowledgeService 管理知识文档和向量检索。', bold=False)
add_para('④ Mapper/Repository 层负责数据持久化。本项目使用 Spring JDBC 直接编写 SQL，Mapper 类封装了 CRUD 操作。', bold=False)
add_para('⑤ Agent Core 通过 LLMClient 接口调用大模型，具体实现可以是 MockLLMClient（离线 Demo）或 OpenAiCompatibleLLMClient（对接 DeepSeek 等）。', bold=False)
add_para('⑥ AgentExtension 提供了工具能力的统一入口：向量检索通过 SemanticVectorSearchExtension 调用 Weaviate，飞书操作通过 LarkCliMcpExtension 调用 lark-cli。新增扩展只需实现 AgentExtension 接口即可。', bold=False)

add_heading('3.2  包图', 2)

add_figure(
    '图 3-2  后端包图',
    '包图展示了后端 12 个 Maven 模块的包命名规范和层次关系。每个模块对应一个 com.flowmind.* 顶级包，'
    '内部按 controller / service / mapper / entity / dto / config / extension 等子包组织。'
    'contrib.capability 作为小组成员能力扩展的缓冲区。',
    '''
com.flowmind
├── com.flowmind.app              ← 聚合启动入口 (FlowMindApplication)
├── com.flowmind.gateway           ← 统一路由 / CORS
├── com.flowmind.user              ← 用户 / 登录 / 角色 / 权限
│   ├── controller
│   ├── service
│   ├── mapper
│   ├── entity
│   ├── dto
│   └── security
├── com.flowmind.agent             ← AI 智能体核心 ★
│   ├── controller                 (AgentController)
│   ├── core                       (Agent / BaseAgent / ContentAgent / ...)
│   ├── service                    (AgentRouter / ConversationService / TraceService)
│   ├── llm                        (LLMClient / MockLLMClient / OpenAiCompatible...)
│   ├── extension                  (AgentExtension / McpToolProvider / SkillProvider / ...)
│   ├── entity / dto / mapper
│   └── config
├── com.flowmind.knowledge         ← 知识库 / 向量检索
│   ├── controller
│   ├── service                    (KnowledgeService)
│   ├── vector                     (EmbeddingService / WeaviateClientService / TextChunker)
│   ├── mapper / entity
│   └── config
├── com.flowmind.content           ← 内容运营 ★
│   ├── controller
│   ├── service                    (ContentLibraryService / ContentSopService)
│   ├── port                       (抽象端口: KnowledgeRetriever / ContentGenerationClient / ...)
│   ├── adapter                    (适配器实现: Mock... / MySql...)
│   ├── entity / dto / mapper / vo
│   └── config
├── com.flowmind.student           ← 学员管理
│   ├── controller / service / mapper / entity
├── com.flowmind.school            ← 院校情报
│   ├── controller / service / mapper / entity
├── com.flowmind.analytics         ← 数据分析
│   ├── controller / service
├── com.flowmind.feishu            ← 飞书生态集成
│   ├── controller / service / FeishuMockClient
├── com.flowmind.common.core       ← 公共: ApiResponse / IdGenerator
├── com.flowmind.common.security   ← 公共: AuthInterceptor / TokenUtil
├── com.flowmind.common.web        ← 公共: OpenApiConfig / WebMvcConfig
└── com.flowmind.contrib.capability ← 扩展缓冲区
    └── vectorsearch               (VectorSearchRuntimeExtension)
''',
    ''
)

add_para(
    '包设计原则说明：'
)
add_para('① contrib.capability 是小组成员扩展能力的缓冲区。团队成员可在该包下新增 ToolService、Controller、Extension，在稳定后再迁移到正式业务模块。这降低了多人协作时的代码冲突。', bold=False)
add_para('② 各包职责不能重叠：Agent 模块负责任务编排，不负责所有业务存储；Knowledge 模块负责知识检索，不负责内容生产；Content 模块负责内容资产，不负责飞书底层实现。', bold=False)
add_para('③ content-service 采用 Port/Adapter（端口/适配器）架构：port 包定义抽象端口（KnowledgeRetriever、ContentGenerationClient 等），adapter 包提供具体实现（Mock 和 MySQL）。这种设计使得替换底层实现（如从 Mock 切换到真实飞书 API）时无需修改 Service 层代码。', bold=False)

add_heading('3.3  前端模块结构', 2)

add_para(
    '前端代码组织如下（基于实际 src/ 目录结构）：'
)

add_code_block('''
frontend/src/
├── api/                  ← API 客户端 (client.ts: fetch + SSE 封装)
├── assets/               ← 静态资源
├── components/           ← 通用组件
├── layouts/              ← 布局组件 (MainLayout.vue)
├── router/               ← Vue Router 路由配置 (index.ts)
├── stores/               ← Pinia 状态管理 (auth.ts 等)
├── views/                ← 页面视图
│   ├── LoginView.vue     ← 登录页
│   ├── SettingsView.vue  ← 设置页
│   └── ...               ← AI 工作台、内容管理、学员管理等
└── App.vue               ← 根组件
''', 'TypeScript/Vue')

add_heading('3.4  后端模块结构', 2)

add_para(
    '后端每个业务模块内部遵循统一的包结构：controller（接收请求）、service（业务逻辑）、mapper（数据访问）、'
    'entity（数据实体）、dto（数据传输对象）、config（模块配置）。以 ai-agent-service 为例：'
)

add_code_block('''
ai-agent-service/src/main/java/com/flowmind/agent/
├── controller/           ← AgentController, PromptController
├── core/                 ← Agent 接口, BaseAgent, ContentAgent, KnowledgeAgent, ...
├── service/              ← AgentRouter, ConversationService, AgentTraceService, ...
├── llm/                  ← LLMClient 接口, MockLLMClient, OpenAiCompatibleLLMClient
├── extension/            ← AgentExtension, McpToolProvider, SkillProvider, ...
├── entity/               ← AgentSessionEntity, ConversationEntity
├── dto/                  ← AgentRequest, AgentResponse
├── mapper/               ← SessionMapper, ConversationMapper
└── config/               ← 模块级配置
''', 'Java')

add_heading('3.5  外部系统适配层设计', 2)

add_para(
    '系统通过适配层（Adapter Layer）与外部系统交互，避免业务逻辑直接依赖外部 SDK。关键适配组件包括：'
)

add_table(
    ['适配组件', '外部系统', '实现方式', '状态'],
    [
        ['OpenAiCompatibleLLMClient', 'LLM API (DeepSeek / OpenAI / 豆包 / 通义千问)', 'HTTP Client + OpenAI 协议适配', '已实现'],
        ['MockLLMClient', '—', '返回预设模拟数据', '已实现'],
        ['LarkCliMcpExtension / LarkCliToolService', '飞书开放平台', '通过 lark-cli 命令行桥接', '部分实现 + Mock 预留'],
        ['WeaviateClientService', 'Weaviate 向量数据库', 'GraphQL 接口', '已实现'],
        ['EmbeddingService', 'Embedding API', 'HTTP 调用嵌入模型', '已实现'],
        ['FeishuMockClient', '—', '返回模拟飞书数据', '已实现'],
        ['WebMvcConfig / CorsFilter', '浏览器跨域', 'Spring CORS 配置', '已实现'],
    ]
)

add_heading('3.6  移动客户端（Android）', 2)

add_para(
    'FlowMind Agent 已实现一版 Android 移动客户端（Debug APK），位于 app/mobile/ 目录。'
    '移动端通过后端 REST API 访问所有业务能力，不直接持有大模型 API Key 或数据库连接。'
)

add_heading('3.6.1  技术栈', 3)
add_table(
    ['技术', '用途'],
    [
        ['Android (Java)', '原生 Android 应用开发，MainActivity.java 约 1,427 行'],
        ['OkHttp / HttpURLConnection', '网络请求与 SSE 流式对话'],
        ['Gradle', '构建系统，app/build.gradle + settings.gradle'],
        ['Android Manifest', '应用权限与组件声明'],
    ]
)

add_heading('3.6.2  核心功能', 3)
add_para(
    '移动端通过 ApiClient.java（约 198 行）统一封装后端接口调用，覆盖以下功能：'
)
add_para('① 用户登录与 Token 管理：与 Web 端共享同一套用户认证体系。')
add_para('② AI 对话工作台：支持流式（SSE）对话，展示模型 Thinking 和工具调用结果。')
add_para('③ 内容运营：选题浏览、文案查看、内容日历。')
add_para('④ 知识库检索：文档列表、标签筛选、向量语义检索。')
add_para('⑤ 学员管理：学员列表、画像查看、AI 分析。')
add_para('⑥ 院校情报：项目浏览、截止时间、报名条件查询。')
add_para('⑦ 系统设置与飞书同步状态查看。')

add_para(
    '移动端的架构原则与 Web 端一致：所有外部 API 调用均经过后端代理，移动端不直接持有 LLM Key 或飞书 App Secret，'
    '敏感操作通过后端鉴权拦截器控制。'
)

add_heading('3.7  桌面客户端（Python / PySide6）', 2)

add_para(
    '系统额外提供了一套 Python 桌面客户端（desktop_fronted/），使用 PySide6 高保真复刻了 Web 前端全部 10 个业务页面。'
    '桌面端采用"保留后端、只重写客户端界面"的迁移策略，通过统一 API 客户端（api.py）访问现有 Java 后端，'
    '不重复开发业务逻辑。'
)

add_heading('3.7.1  技术栈', 3)
add_table(
    ['技术', '版本', '用途'],
    [
        ['PySide6', '6.7.3', '高保真 GUI 框架，复刻 Web 端页面和交互'],
        ['Tkinter', '标准库', '兜底 GUI 方案，无需额外安装依赖即可运行'],
        ['httpx', '0.27+', 'HTTP/SSE 网络请求'],
        ['Python', '3.10+', '开发语言'],
    ]
)

add_heading('3.7.2  架构设计', 3)
add_para(
    '桌面端由以下核心模块组成：'
)
add_para('① api.py（~36KB）：统一 API 客户端，封装所有后端 REST 接口和 SSE 流式端点，同时内置离线 Mock 模式（offline://demo），无后端时可完整演示。')
add_para('② fallback_data.py（~16KB）：离线 Demo 数据，覆盖所有业务模块的 Mock 数据，支持模拟登录、角色权限、流式对话、内容运营、知识库、学员/院校等。')
add_para('③ views.py（~224KB）：PySide6 主页面实现，包含 LoginPage、ShellWidget（主布局）、DashboardPage、AgentPage（AI 工作台）、ContentPage、KnowledgePage、StudentsPage、SchoolsPage、FeishuPage、SettingsPage 共 10 个业务页面。')
add_para('④ widgets.py（~21KB）：通用组件库，包含 Badge、StarRating、MarkdownPanel、JsonPanel、TraceListPanel、CalendarGrid、ChartCard、DataTable、StreamWorker（SSE 流式线程）、ApiWorker（异步 API 线程）等。')
add_para('⑤ styles.py（~6KB）：QSS 样式系统，定义全应用统一的颜色、字体、圆角、边框、hover/selected 状态等视觉规范。')

add_heading('3.7.3  与 Web 端的功能对齐程度', 3)
add_table(
    ['页面', 'Web 端', '桌面端 (PySide6)', '对齐度'],
    [
        ['登录', 'LoginView.vue', 'LoginPage（双栏品牌面板 + 5 个 Demo 账号快捷入口）', '✅ 完整对齐'],
        ['Dashboard', 'DashboardView.vue', 'DashboardPage（Hero 区 + 统计卡片 + 图表）', '✅ 完整对齐'],
        ['AI 工作台', 'AgentWorkspaceView.vue', 'AgentPage（会话管理 + SSE 流式 + Thinking/Trace + 上下文面板）', '✅ 完整对齐'],
        ['内容运营', 'ContentView.vue', 'ContentPage（主题库 + 文案库 + 日历 + SOP + 一键入库）', '✅ 完整对齐'],
        ['知识库', 'KnowledgeView.vue', 'KnowledgePage（卡片/表格视图 + 标签编辑 + 向量检索 + 同步）', '✅ 完整对齐'],
        ['学员管理', 'StudentsView.vue', 'StudentsPage（卡片/表格 + 风险着色 + AI 分析 + 进度条）', '✅ 完整对齐'],
        ['院校情报', 'SchoolsView.vue', 'SchoolsPage（项目卡片 + 搜索筛选 + 截止着色 + AI 推荐）', '✅ 完整对齐'],
        ['数据分析', 'AnalyticsView.vue', 'AnalyticsPage（统计卡 + 自绘图表）', '✅ 完整对齐'],
        ['飞书同步', 'FeishuView.vue', 'FeishuPage（连接状态 + 同步日志 + 文档创建/读取）', '✅ 完整对齐'],
        ['系统设置', 'SettingsView.vue', 'SettingsPage（模型/Prompt/飞书/权限/路由/日志）', '✅ 完整对齐'],
    ]
)

add_para(
    '桌面端还支持以下特色能力：'
)
add_para('① 离线 Demo 模式：使用 offline://demo 作为 Base URL 时，完全离线运行，内置 Mock 数据覆盖所有业务场景。')
add_para('② 双渲染器：PySide6 作为主渲染器提供高保真体验，Tkinter（tk_app.py）作为兜底渲染器在 PySide6 不可用时仍可运行。')
add_para('③ 自动化自检：smoke_test.py（API/离线数据）、pyside_smoke_test.py（GUI 页面创建/切换）、check_desktop.py（一键全检）。')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第四章 数据架构设计
# ═══════════════════════════════════════════════
add_heading('4  数据架构设计', 1)

add_heading('4.1  数据存储总体方案', 2)

add_para(
    '系统采用分层存储架构：MySQL 作为业务事实源存储所有结构化业务数据，Weaviate 作为向量检索引擎存储知识文档的'
    '向量化分片，Redis 预留用于缓存和会话管理，MinIO 预留用于图片和文件对象存储。'
)

add_table(
    ['存储系统', '定位', '存储内容', '状态'],
    [
        ['MySQL', '业务事实源（Source of Truth）', '用户/角色/权限、AI 会话/消息、内容主题/文案/日历、知识文档/标签、学员/院校、飞书同步记录、系统日志', '已实现（17 张表）'],
        ['Weaviate', '语义检索引擎', '知识文档的向量化分片（chunk），关联 MySQL 中的原始文档', '已实现'],
        ['Redis', '缓存与中间层', '热点配置缓存、会话 Token、API 限流计数器、统计分析缓存', '预留'],
        ['MinIO', '对象存储', '内容配图、文件附件、学员头像、知识库文件', '预留'],
    ]
)

add_heading('4.2  MySQL 业务数据设计', 2)

add_para(
    'MySQL 数据库名为 FlowMind，字符集使用 utf8mb4。按照业务域划分，共有 19 张核心表。每张表均包含 id（BIGINT 主键自增）、'
    'created_at、updated_at、deleted 字段，支持逻辑删除和时间审计。'
)

add_para('各业务域表结构如下：', bold=True)

add_heading('4.2.1  用户权限域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['sys_user', '用户表', 'username, password, nickname, avatar, status'],
        ['sys_role', '角色表', 'role_code, role_name'],
        ['sys_user_role', '用户-角色关联表', 'user_id, role_id'],
        ['sys_permission', '权限表', 'permission_code, permission_name, path_pattern, frontend_route'],
        ['sys_role_permission', '角色-权限关联表', 'role_id, permission_id'],
    ]
)

add_heading('4.2.2  AI 会话域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['agent_session', 'Agent 会话表', 'user_id, agent_type, title'],
        ['agent_message', 'Agent 消息表', 'session_id, role, content, cards (JSON)'],
        ['prompt_template', 'Prompt 模板表', 'agent_type, name, template, enabled'],
    ]
)
add_para(
    'agent_message 的 metadata/cards 字段以 JSON 格式存储工具调用 Trace、模型 Thinking/Reasoning 过程、'
    '处理状态等扩展信息。metadata 字段设计为可空，兼容旧数据的向后兼容。'
)

add_heading('4.2.3  内容运营域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['content_topic', '内容主题库', 'title, platform, topic_type, style, status, heat_score, favorite'],
        ['content_calendar', '内容日历', 'topic_id, publish_date, channel, owner, status'],
        ['content_copy', '文案库', '（多版本文案、评分、使用状态、使用日期）'],
        ['content_copy_image', '配图表', '（文案配图/建议图片管理）'],
        ['content_generation_record', '内容生成流水', '（输入输出摘要、生成类型）'],
        ['content_publish_metric', '发布效果指标', '（平台发布后的效果数据）'],
    ]
)

add_heading('4.2.4  知识库域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['knowledge_doc', '知识文档表', 'title, category, file_url, summary, source'],
        ['knowledge_tag', '知识标签表', 'name, color'],
        ['knowledge_doc_tag', '文档-标签关联表', 'doc_id, tag_id'],
        ['knowledge_sync_log', '知识同步日志表', '（飞书文档同步记录）'],
    ]
)

add_heading('4.2.5  学员域与院校域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['student_profile', '学员画像表', 'name, gender, school, major, gpa, ranking, english_score, target_school, risk_level'],
        ['student_application_progress', '申请进度表', 'student_id, stage, progress, note'],
        ['school_info', '院校信息表', 'name, region, level, discipline_tags'],
        ['school_project', '院校项目表', 'school_id, project_name, project_type, deadline, requirements, materials'],
    ]
)

add_heading('4.2.6  飞书与日志域', 3)
add_table(
    ['表名', '说明', '关键字段'],
    [
        ['feishu_sync_record', '飞书同步记录表', 'sync_type, target_name, status, message'],
        ['system_log', '系统日志表', 'module, level, message, trace_id'],
    ]
)

add_heading('4.3  Weaviate 向量数据设计', 2)

add_para(
    'Weaviate 中定义 KnowledgeDoc 类（class），用于存储知识文档的向量化分片。每个分片（chunk）对应一条 Weaviate 对象。'
)

add_table(
    ['字段', '类型', '说明'],
    [
        ['title', 'text', '文档标题'],
        ['chunkText', 'text', '分片文本内容'],
        ['feishuToken', 'text', '飞书文档 token（用于关联飞书源文档）'],
        ['feishuUrl', 'text', '飞书文档 URL'],
        ['feishuType', 'text', '飞书文档类型（doc/sheet/bitable/folder）'],
        ['tags', 'text[]', '标签数组'],
        ['mysqlId', 'int', '关联 MySQL knowledge_doc 表的 id'],
        ['chunkIndex', 'int', '分片序号（第几个 chunk）'],
        ['totalChunks', 'int', '文档总分片数'],
    ]
)

add_heading('4.4  Redis 与 MinIO 预留设计', 2)

add_para(
    'Redis 预留了以下缓存键模式：'
)
add_para('• cache:config:{key} — 热点系统配置缓存', bold=False)
add_para('• cache:knowledge:stats — 知识库统计数据', bold=False)
add_para('• session:{token} — 用户会话 Token', bold=False)
add_para('• rate-limit:{userId}:{api} — API 限流计数器', bold=False)

add_para(
    'MinIO 预留了以下存储桶（Bucket）：'
)
add_para('• content-images — 内容运营配图', bold=False)
add_para('• knowledge-files — 知识库附件文件', bold=False)
add_para('• student-avatars — 学员头像', bold=False)

add_heading('4.5  数据一致性与索引重建策略', 2)

add_para(
    '数据一致性设计遵循以下原则：'
)
add_para('① MySQL 是业务事实源（Source of Truth）：所有业务写入必须先成功写入 MySQL，再异步更新 Weaviate 索引。', bold=False)
add_para('② Weaviate 是检索索引，不替代 MySQL：即使 Weaviate 完全宕机，系统仍可通过 MySQL 的 LIKE 搜索提供降级服务。', bold=False)
add_para('③ 飞书 token 与本地知识文档保留映射关系：knowledge_doc 表中保存飞书文档 token，确保可以追溯原始来源。', bold=False)
add_para('④ 向量库重建策略：当 Weaviate 数据丢失或 schema 变更时，根据 MySQL 中"未进入向量库"或"已变更"的知识文档重新执行 Embedding 和写入操作。重建通过 knowledge_sync_log 表追踪进度。', bold=False)
add_para('⑤ 会话消息 metadata 字段兼容设计：agent_message 的 JSON 字段可为空，旧数据没有 Thinking/Trace 信息时前端应正常展示，不报错。', bold=False)

add_heading('4.6  ER 图', 2)

add_figure(
    '图 4-1  核心数据实体关系图（简化版）',
    '该图展示了 MySQL 核心表之间的主要关联关系。实际表结构包含更多字段，此处仅展示关键外键关系。',
    '''
┌──────────────┐     ┌──────────────────┐     ┌─────────────────────┐
│   sys_user   │────→│   sys_user_role  │←────│      sys_role       │
│              │     └──────────────────┘     │                     │
│   id (PK)    │                              │   id (PK)            │
│   username   │                              │   role_code          │
│   password   │     ┌──────────────────┐     │   role_name          │
│   nickname   │     │ sys_role_permission│←───│                     │
└──────┬───────┘     └────────┬─────────┘     └─────────────────────┘
       │                      │
       │                      │              ┌──────────────────────┐
       │                      └──────────────│   sys_permission     │
       │                                     │   id (PK)            │
       ▼                                     │   permission_code    │
┌──────────────┐                             │   path_pattern       │
│agent_session │                             └──────────────────────┘
│   id (PK)    │
│   user_id(FK)│
│   agent_type │
│   title      │
└──────┬───────┘
       │
       ▼
┌──────────────┐     ┌──────────────────┐     ┌──────────────────────┐
│agent_message │     │  content_topic   │────→│  content_calendar    │
│   id (PK)    │     │   id (PK)        │     │   id (PK)            │
│   session_id │     │   title          │     │   topic_id (FK)      │
│   role       │     │   platform       │     │   publish_date       │
│   content    │     │   status         │     └──────────────────────┘
│   cards(JSON)│     └──────────────────┘
└──────────────┘
                   ┌──────────────────┐     ┌──────────────────────┐
                   │ knowledge_doc    │────→│ knowledge_doc_tag    │←──┐
                   │   id (PK)        │     └──────────────────────┘   │
                   │   title          │                                │
                   │   summary        │     ┌──────────────────────┐   │
                   │   source         │     │  knowledge_tag       │───┘
                   └──────────────────┘     │   id (PK) / name     │
                                            └──────────────────────┘
┌─────────────────┐     ┌───────────────────────────┐
│ student_profile │────→│student_application_progress│
│   id (PK)       │     │   id (PK)                 │
│   name          │     │   student_id (FK)         │
│   gpa           │     │   stage                   │
│   risk_level    │     │   progress                │
└─────────────────┘     └───────────────────────────┘

┌──────────────────┐    ┌───────────────────────┐
│   school_info    │───→│   school_project      │
│   id (PK)        │    │   id (PK)             │
│   name           │    │   school_id (FK)      │
│   region         │    │   project_name        │
│   level          │    │   deadline            │
└──────────────────┘    └───────────────────────┘
''',
    ''
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第五章 服务协作与事务设计
# ═══════════════════════════════════════════════
add_heading('5  服务协作与事务设计', 1)

add_heading('5.1  AI 工作台问答协作流程', 2)

add_para(
    '这是系统最核心的协作流程——用户在前端 AI 工作台输入问题，系统自动路由到合适的智能体，调用大模型和工具后流式返回回答。'
)

add_para('详细步骤：', bold=True)
add_para('① 前端通过 EventSource 发起 SSE 请求，URL 携带 agent_type 参数（或 "auto" 交给系统自动判断）。', bold=False)
add_para('② AgentController 接收请求，先调用 ConversationService 保存用户消息（role="user"）到 agent_message 表。', bold=False)
add_para('③ AgentRouter 根据用户输入内容进行关键词匹配，自动路由到对应的 Agent（ContentAgent / KnowledgeAgent / StudentAgent / SchoolAgent / FeishuAgent）。如果输入不含明确领域关键词，默认路由到 ContentAgent。', bold=False)
add_para('④ AgentTraceService 初始化工具调用上下文，收集所有可用的 AgentExtension 的 runtimeContext() 信息。', bold=False)
add_para('⑤ AgentExtension 链依次执行：RuntimeToolExtensions 提供时间和联网信息；SemanticVectorSearchExtension 执行向量检索并将结果注入上下文；LarkCliMcpExtension 执行飞书相关操作。', bold=False)
add_para('⑥ LLMClient.stream() 被调用，系统 Prompt 包含 Agent 角色定义、工具上下文和对话历史。大模型以流式方式返回 token。', bold=False)
add_para('⑦ 后端解析流式响应中的 reasoning_content（模型 Thinking 过程）和 content（最终回答），通过 SSE 分别以 reasoning、thinking、delta、trace、done 等事件类型推送给前端。', bold=False)
add_para('⑧ 前端实时展示流式回答：delta 事件追加到回答区，thinking 事件追加到折叠的思考过程区，trace 事件展示工具调用结果。', bold=False)
add_para('⑨ 流式结束后，后端将 assistant 消息（含完整的 Thinking、Trace 和回答）保存到 agent_message 表。', bold=False)

add_para(
    '此流程的同步特征：用户问答是同步请求 + 流式响应，工具调用在模型生成前完成作为上下文注入。LLM 生成过程通过 SSE 实时返回。'
)

add_heading('5.2  内容生成与入库协作流程', 2)

add_para(
    '当用户在 AI 工作台触发内容 SOP（如"生成小红书文案"），系统执行以下协作流程：'
)
add_para('① AgentRouter 将请求路由到 ContentAgent。', bold=False)
add_para('② ContentAgent 通过 AgentExtension 收集知识库中的爆款结构参考（向量检索相似文案）、内容主题库中的历史选题。', bold=False)
add_para('③ ContentSopService 根据平台（小红书/朋友圈/公众号）选择对应的 SOP 流程和 Prompt 模板。', bold=False)
add_para('④ LLMClient 调用大模型生成文案初稿，SSE 流式返回给前端。', bold=False)
add_para('⑤ 用户确认后，ContentLibraryService 保存主题到 content_topic、文案到 content_copy、配图到 content_copy_image。', bold=False)
add_para('⑥ content_calendar 记录计划发布日期和渠道。', bold=False)
add_para('⑦ 可选：通过 FeishuAgent 将文案同步到飞书文档或多维表格。', bold=False)
add_para('⑧ content_generation_record 保存本次生成的输入摘要和输出摘要。', bold=False)

add_heading('5.3  知识库同步与向量检索协作流程', 2)

add_para(
    '知识库文档从飞书同步并建立向量索引的流程：'
)
add_para('① feishu-service 或 knowledge-service 通过 lark-cli 读取飞书共享文件夹中的文档列表。', bold=False)
add_para('② knowledge-service 将文档信息（标题、摘要、来源、飞书 token）保存到 knowledge_doc 表。', bold=False)
add_para('③ TextChunker 将文档内容按段落切分为 chunk（每 chunk 约 500-1000 字符）。', bold=False)
add_para('④ EmbeddingService 调用嵌入模型 API 为每个 chunk 生成向量。', bold=False)
add_para('⑤ WeaviateClientService 将向量 + 元数据写入 Weaviate 的 KnowledgeDoc class。', bold=False)
add_para('⑥ 查询时：VectorSearchToolService 接收查询文本 → EmbeddingService 生成查询向量 → Weaviate 近邻搜索返回 TopK 结果。', bold=False)
add_para('⑦ 若 Weaviate 不可用或无结果，fallback 到 MySQL 的 LIKE 搜索作为降级方案。', bold=False)

add_heading('5.4  跨服务事务与最终一致性设计', 2)

add_para(
    'Demo 阶段以单体聚合方式运行，所有服务共享同一个 JVM 和数据库连接，可以依赖 Spring 的本地事务管理（@Transactional）。'
    '当未来拆分为真实微服务后，跨服务操作不应使用分布式强事务（如 2PC），而应采用最终一致性方案：'
)
add_para('① 本地事务先保存主业务数据（如内容文案写入 MySQL）。', bold=False)
add_para('② 生成同步任务或事件（如"同步到飞书文档"、"更新向量索引"）。', bold=False)
add_para('③ 下游服务异步处理：飞书同步服务监听事件后调用飞书 API，向量更新服务监听事件后重建向量索引。', bold=False)
add_para('④ feishu_sync_record 表保存每次同步的状态和失败原因，失败支持重试。', bold=False)
add_para('⑤ 对外部系统（如飞书 API）的调用不应和本地数据库操作强绑定一个事务。应保存外部调用结果和失败原因，保证可追溯。', bold=False)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第六章 认证鉴权与安全设计
# ═══════════════════════════════════════════════
add_heading('6  认证、鉴权与安全设计', 1)

add_heading('6.1  当前 Demo 鉴权方案', 2)

add_para(
    'Demo 阶段实现了基本的认证鉴权流程：'
)
add_para('① 用户通过 /api/auth/login 接口提交用户名和密码，后端验证后返回 Token。', bold=False)
add_para('② 前端 Pinia auth store 保存 Token，后续请求通过 Authorization: Bearer <token> Header 携带。', bold=False)
add_para('③ Vue Router 的全局前置守卫（beforeEach）检查 token 是否存在：无 token 跳转登录页，有 token 放行。', bold=False)
add_para('④ AuthInterceptor（位于 common-security 模块）在后端拦截请求，校验 Token 有效性。', bold=False)
add_para('⑤ TokenUtil 负责 Token 的生成和解析。', bold=False)

add_heading('6.2  后续 JWT 与 RBAC 方案', 2)

add_para(
    'Demo 阶段已预留了完整的 RBAC 表结构（sys_user、sys_role、sys_permission、sys_user_role、sys_role_permission），'
    '后续可直接升级为 JWT + RBAC 方案：'
)
add_para('① Token 从简单 Token 升级为 JWT，包含用户 ID、角色列表、过期时间等信息。', bold=False)
add_para('② SecurityPermissionService（已实现）负责查询用户的所有权限路径。', bold=False)
add_para('③ 接口路径通过 sys_permission.path_pattern 与前端路由 sys_permission.frontend_route 绑定。', bold=False)
add_para('④ 角色设计：管理员（admin）、运营人员（operator）、咨询老师（consultant）、只读用户（viewer）。', bold=False)
add_para('⑤ 敏感操作（如删除用户、修改角色权限、删除知识文档）需要管理员角色。', bold=False)
add_para('⑥ RbacDataInitializer（已实现）在应用启动时初始化默认角色和权限数据。', bold=False)

add_heading('6.3  敏感配置管理', 2)

add_para(
    '系统严格遵循"敏感信息不暴露"原则：'
)
add_para('① LLM API Key、飞书 App Secret、数据库密码等不放入公共配置文件。', bold=False)
add_para('② 使用 application-local.yml 保存本地私有配置，该文件已加入 .gitignore，不进入版本控制。', bold=False)
add_para('③ 前端不持有 LLM Key 或飞书 Key——所有外部 API 调用均经过后端代理。', bold=False)
add_para('④ 生产环境建议使用环境变量（${ENV_VAR}）或配置中心注入敏感值。', bold=False)

add_heading('6.4  工具调用安全边界', 2)

add_para(
    'AI Agent 的工具调用需要安全边界控制：'
)
add_para('① 工具调用白名单：Agent 只能调用已注册的 AgentExtension，不能执行任意代码或命令。', bold=False)
add_para('② 飞书工具需要用户授权：飞书的写入操作（创建文档、修改多维表格）需要有效的飞书 Token，Token 过期或权限不足时明确提示用户。', bold=False)
add_para('③ Shell 命令执行不暴露给 LLM：系统不提供任意 Shell 执行扩展，避免 Prompt 注入导致安全问题。', bold=False)
add_para('④ 外部 URL 请求需要白名单校验：RuntimeToolExtensions 中的联网能力需要限制可访问的域名范围。', bold=False)

add_heading('6.5  审计与日志安全', 2)

add_para(
    '所有敏感操作记录到 system_log 表，包括：登录/登出、角色权限变更、飞书同步操作、知识库删除、学员数据修改。'
    '日志中包含 module（模块）、level（等级）、trace_id（追踪 ID），便于安全审计和故障回溯。'
)

add_figure(
    '图 6-1  认证鉴权流程',
    '展示了用户从登录到访问受保护接口的完整鉴权流程。',
    '''
用户          前端(Vue)        后端(Spring)       数据库
 │               │                │                 │
 │  输入账号密码  │                │                 │
 │──────────────→│                │                 │
 │               │  POST /login   │                 │
 │               │───────────────→│                 │
 │               │                │  查询用户        │
 │               │                │────────────────→│
 │               │                │←────────────────│
 │               │                │  生成Token       │
 │               │  {token,user}  │                 │
 │               │←───────────────│                 │
 │               │  存Pinia Store │                 │
 │               │                │                 │
 │  访问受保护页  │                │                 │
 │──────────────→│                │                 │
 │               │  路由守卫检查token               │
 │               │  GET /api/xxx  │                 │
 │               │  Authorization:│                 │
 │               │  Bearer <token>│                 │
 │               │───────────────→│                 │
 │               │                │ AuthInterceptor │
 │               │                │  校验Token       │
 │               │                │  查询权限        │
 │               │                │────────────────→│
 │               │                │←────────────────│
 │               │  200 {data}    │                 │
 │               │←───────────────│                 │
 │               │                │                 │
''',
    ''
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第七章 高并发与性能设计
# ═══════════════════════════════════════════════
add_heading('7  高并发与性能设计', 1)

add_heading('7.1  潜在性能瓶颈', 2)

add_para(
    '虽然本系统为课程 Demo，不承载大规模生产流量，但应识别潜在的并发与性能压力点，以便后续扩展时有所准备：'
)

add_table(
    ['压力点', '场景描述', '影响范围'],
    [
        ['多用户同时使用 AI 工作台', '多个用户同时发起 SSE 流式请求，每个请求持有数据库连接和 HTTP 连接', '后端并发连接数、数据库连接池'],
        ['LLM API 响应慢', '大模型 API 单次调用可能耗时 5-30 秒', '用户体验、SSE 超时设置'],
        ['向量检索耗时', 'Weaviate 查询在大量文档时可能变慢', 'AI 回答等待时间'],
        ['飞书 API 限流', '飞书开放平台对调用频率有限制（通常 100 次/秒/应用）', '飞书同步成功率'],
        ['SSE 长连接占用资源', '每个 SSE 连接占用一个 HTTP 线程直到流式响应结束', 'Tomcat 线程池'],
        ['数据分析图表频繁刷新', '统计查询涉及多表 JOIN 和聚合', 'MySQL 查询耗时'],
    ]
)

add_heading('7.2  前端优化', 2)
add_para('① 列表分页和懒加载：学员列表、内容主题列表、知识文档列表均使用后端分页，前端仅加载当前页数据。', bold=False)
add_para('② 防抖节流：搜索输入框使用 300ms 防抖，避免每次按键触发后端查询。', bold=False)
add_para('③ ECharts 按需加载：图表仅在对应标签页激活时才初始化和请求数据，减少不必要的网络请求。', bold=False)
add_para('④ SSE 连接管理：页面离开时关闭 EventSource 连接，避免后台消耗资源。', bold=False)

add_heading('7.3  后端优化', 2)
add_para('① 数据库连接池：使用 HikariCP（Spring Boot 默认连接池），配置合理的最大连接数和超时时间。', bold=False)
add_para('② LLM 请求超时控制：设置 60 秒超时，超时后向用户返回降级提示而非无限等待。', bold=False)
add_para('③ 工具调用超时和降级：每个 AgentExtension 设置独立超时（如向量检索 5 秒、飞书操作 10 秒），失败时降解为不包含该工具结果的回答。', bold=False)
add_para('④ SSE 连接超时管理：设置合理的 SSE 超时时间，避免僵尸连接占用线程。', bold=False)

add_heading('7.4  数据库与缓存优化', 2)
add_para('① MySQL 索引：在 user_id、session_id、platform、created_at 等高频查询字段上建立索引。', bold=False)
add_para('② 分页查询：所有列表接口强制使用 LIMIT + OFFSET 分页，默认每页 20 条。', bold=False)
add_para('③ Redis 缓存（预留）：热点配置、知识库统计、常用查询结果放入 Redis 缓存，减少 MySQL 压力。', bold=False)
add_para('④ 向量检索 topK 限制：默认 topK=5，避免返回过多无关结果。', bold=False)

add_heading('7.5  外部 API 限流与降级', 2)
add_para('① 飞书 API 调用限流和重试：调用频率超过限制时指数退避重试（最多 3 次）。', bold=False)
add_para('② LLM API 熔断：连续失败 N 次后进入熔断状态，后续请求直接返回 MockLLMClient 的降级回答。', bold=False)
add_para('③ Embedding API 降级：Embedding 服务不可用时，向量检索降级为 MySQL LIKE 文本搜索。', bold=False)

add_heading('7.6  后续扩展方案', 2)
add_para(
    '当系统需要承载更大规模时：可将 ai-agent-service 和 knowledge-service 水平扩展为多实例，'
    '通过 Nginx 或 Spring Cloud Gateway 做负载均衡；引入 Redis 集中管理会话 Token，使各实例无状态化；'
    '引入消息队列（RabbitMQ/Kafka）处理飞书同步、向量更新等异步任务；'
    '引入 Prometheus + Grafana 监控各服务实例的 QPS、延迟和错误率。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第八章 日志监控与故障定位
# ═══════════════════════════════════════════════
add_heading('8  日志、监控与故障定位', 1)

add_heading('8.1  日志类型', 2)

add_table(
    ['日志类型', '记录内容', '存储位置'],
    [
        ['系统运行日志', 'Spring Boot 应用启动、配置加载、Bean 注册、异常堆栈', '控制台 + 文件日志'],
        ['用户登录日志', '登录成功/失败、IP 地址、时间', 'system_log 表'],
        ['AI 会话日志', '每次对话的 session_id、agent_type、输入输出长度', 'agent_message 表 + 应用日志'],
        ['工具调用 Trace', '工具名称、输入参数、输出结果、耗时', 'agent_message.metadata (JSON)'],
        ['模型 Thinking/Reasoning', '大模型返回的思考过程文本', 'agent_message.metadata (JSON)'],
        ['飞书同步日志', '同步类型、目标名称、状态（成功/失败）、错误信息', 'feishu_sync_record 表'],
        ['知识库同步日志', '文档标题、同步状态、向量写入结果', 'knowledge_sync_log 表'],
        ['向量检索日志', '查询文本、返回结果数、distance 值', '应用日志'],
        ['系统操作日志', '角色变更、权限修改、数据删除等敏感操作', 'system_log 表'],
    ]
)

add_heading('8.2  Trace 与 Thinking 的可追溯设计', 2)

add_para(
    '系统在 agent_message 表的 metadata 字段（JSON 格式）中保存了每次 AI 对话的完整可追溯信息。'
    '这类似于 LangChain 的 Callback 机制，但轻量很多。具体包含：'
)
add_para('① tools_called: [{name, input, output, duration_ms}] — 本次对话调用了哪些工具，每个工具的输入输出和耗时。', bold=False)
add_para('② thinking: "..." — 模型的完整思考过程（reasoning_content）。', bold=False)
add_para('③ model: "deepseek-chat" — 使用的模型名称。', bold=False)
add_para('④ prompt_tokens / completion_tokens — Token 用量统计。', bold=False)

add_para(
    '前端在 AI 工作台中可以展开折叠的 "思考过程" 和 "工具调用" 区域，用户可以清晰地看到系统在生成回答前调用了哪些工具、'
    '每个工具返回了什么结果、模型是如何思考的。这大大提升了 AI 系统的透明度和可信度。'
)

add_heading('8.3  故障定位流程', 2)

add_para('以用户反馈"某次回答不准确"为例，故障定位流程如下：')
add_para('① 用户提供会话 ID 或大致时间。', bold=False)
add_para('② 运维人员根据 session_id 查询 agent_message 表，找到对应会话的所有消息。', bold=False)
add_para('③ 查看该消息的 metadata 字段中的工具 Trace：是否调用了向量检索？检索返回了什么结果？是否调用了飞书工具？', bold=False)
add_para('④ 查看模型 Thinking：模型的推理过程是否合理？', bold=False)
add_para('⑤ 判断问题根因：向量检索未命中（知识库缺少相关资料）？飞书工具失败（Token 过期/权限不足）？LLM 生成偏差（Prompt 不够精确）？', bold=False)
add_para('⑥ 修复：补充知识库资料并重建向量索引 / 调整 Prompt 模板 / 修复飞书授权 / 增加降级提示文本。', bold=False)

add_heading('8.4  后续监控平台方案', 2)

add_para(
    'Demo 阶段使用控制台日志和数据库日志表满足基本需求。后续生产部署建议引入：'
)
add_table(
    ['组件', '用途'],
    [
        ['Spring Boot Actuator', '应用健康检查、Metrics 暴露、环境信息查看'],
        ['Prometheus', 'Metrics 采集和存储（QPS、延迟、错误率、JVM 指标）'],
        ['Grafana', '可视化仪表盘（API 调用量、LLM 响应时间、SSE 连接数、Weaviate 查询耗时）'],
        ['ELK / Loki', '集中式日志收集、搜索和分析'],
        ['告警规则', 'LLM API 调用成功率 < 95%、Weaviate 查询平均耗时 > 3s、SSE 连接数超过阈值'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第九章 部署与运维设计
# ═══════════════════════════════════════════════
add_heading('9  部署与运维设计', 1)

add_heading('9.1  当前 Demo 部署图', 2)

add_figure(
    '图 9-1  Demo 本地部署图',
    '当前课程 Demo 阶段所有组件部署在一台开发者电脑上。',
    '''
┌──────────────────────────────────────────────────────────────────┐
│                  开发者电脑 (Windows 11 / macOS)                  │
│                                                                  │
│  ┌─────────────────┐    ┌──────────────────────────────────┐    │
│  │  Browser        │    │  Vite Dev Server (:5173)         │    │
│  │  (Chrome/Edge)  │◄───│  Vue 3 Frontend                  │    │
│  └─────────────────┘    └────────────┬─────────────────────┘    │
│                                      │ REST :8080 + SSE          │
│  ┌───────────────────────────────────┴─────────────────────┐    │
│  │  Spring Boot app-service (:8080)                        │    │
│  │  ┌─────────────────────────────────────────────────┐    │    │
│  │  │  gateway / user / agent / content / knowledge   │    │    │
│  │  │  student / school / analytics / feishu / common │    │    │
│  │  └─────────────────────────────────────────────────┘    │    │
│  └───────┬──────────────┬──────────────┬──────────────────┘    │
│          │              │              │                        │
│  ┌───────┴──┐  ┌────────┴────┐  ┌─────┴──────────┐            │
│  │ MySQL    │  │  Weaviate   │  │  lark-cli      │            │
│  │ Docker   │  │  Docker     │  │  (本地授权)    │            │
│  │ :3306    │  │  :8081      │  │                │            │
│  └──────────┘  └─────────────┘  └────────────────┘            │
│                                                                  │
│  配置文件: application-local.yml (不进入 Git)                    │
└──────────────────────────────────────────────────────────────────┘
''',
    ''
)

add_heading('9.2  本地部署流程', 2)

add_para('当前 Demo 的启动步骤：', bold=True)
add_para('① 拉取 Git 仓库：git clone <repo-url> && cd FlowMind-Agent', bold=False)
add_para('② 启动 MySQL：docker run -d --name mysql -p 3306:3306 -e MYSQL_ROOT_PASSWORD=xxx mysql:9', bold=False)
add_para('③ 启动 Weaviate：docker run -d --name weaviate -p 8081:8080 semitechnologies/weaviate', bold=False)
add_para('④ 配置 application-local.yml：填入本地数据库密码、LLM API Key（如有）、飞书 Token（如有）。无 Key 时系统自动使用 MockLLMClient。', bold=False)
add_para('⑤ 启动后端：cd backend && mvnw spring-boot:run', bold=False)
add_para('⑥ 启动前端：cd frontend && npm install && npm run dev', bold=False)
add_para('⑦ 浏览器访问 http://localhost:5173，使用 Demo 账号登录。', bold=False)

add_heading('9.3  生产部署建议', 2)

add_figure(
    '图 9-2  生产部署图（建议）',
    '后续生产环境的推荐部署架构。前端构建为静态资源部署到 Nginx，后端服务 Docker 化后部署到应用服务器，'
    '数据服务独立部署。通过 Nginx 暴露 HTTPS，使用环境变量注入配置。',
    '''
                          ┌──────────────────────┐
                          │    用户 (HTTPS)       │
                          └──────────┬───────────┘
                                     │
                                     ▼
                          ┌──────────────────────┐
                          │  Nginx / API Gateway │
                          │  • 静态前端资源       │
                          │  • SSL 终止          │
                          │  • 反向代理          │
                          └──────────┬───────────┘
                                     │
          ┌──────────────────────────┼──────────────────────────┐
          │                          │                          │
          ▼                          ▼                          ▼
┌──────────────────┐   ┌──────────────────────┐   ┌──────────────────┐
│  应用服务器 1     │   │   应用服务器 2        │   │  数据服务器      │
│                  │   │                      │   │                  │
│  gateway-service │   │  knowledge-service   │   │  MySQL (主)      │
│  user-service    │   │  content-service     │   │  MySQL (从)      │
│  ai-agent-service│   │  student-service     │   │  Redis           │
│                  │   │  school-service      │   │  Weaviate        │
│                  │   │  analytics-service   │   │  MinIO           │
│                  │   │  feishu-service      │   │                  │
└──────────────────┘   └──────────────────────┘   └──────────────────┘
          │                          │                          │
          └──────────────────────────┼──────────────────────────┘
                                     │
                          ┌──────────┴───────────┐
                          │   External Cloud     │
                          │  • LLM API           │
                          │  • Embedding API     │
                          │  • Feishu Platform   │
                          └──────────────────────┘
''',
    ''
)

add_heading('9.4  配置管理', 2)

add_para(
    'Demo 阶段配置管理策略：'
)
add_para('① application.yml：公共默认配置，进入 Git。包含端口、数据源类型、Weaviate 地址等非敏感信息。', bold=False)
add_para('② application-local.yml：本地私有配置，不进入 Git。包含数据库密码、LLM API Key、飞书 App Secret。', bold=False)
add_para('③ 生产环境：敏感配置通过环境变量（${DB_PASSWORD}、${LLM_API_KEY}）或配置中心注入。', bold=False)

add_heading('9.5  版本升级与演化', 2)

add_para(
    '系统的版本升级和演化策略：'
)
add_para('① 数据库升级：使用 Flyway 或 Liquibase 管理 SQL 迁移脚本，每次变更记录版本号。当前 Demo 阶段直接通过 schema.sql 初始化。', bold=False)
add_para('② 向量库 schema 变更：Weaviate class 新增字段时保持向后兼容，旧字段保留并逐步废弃；重大变更时重建全部向量索引。', bold=False)
add_para('③ 模块演化为微服务：每个模块独立编写启动类和配置文件，通过接口抽象（Feign 接口 + 实现）替换模块间的 Bean 注入。', bold=False)
add_para('④ 各组件演化时通过接口抽象降低影响：如从 MockLLMClient 切换到 OpenAiCompatibleLLMClient，从 Mock 飞书切换到真实飞书 API，业务代码无需重构。', bold=False)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第十章 技术关键点解决方案
# ═══════════════════════════════════════════════
add_heading('10  技术关键点解决方案', 1)

add_para(
    '本章选取 8 个技术关键点，从背景问题、设计目标、方案设计、关键流程、优点、局限与后续优化六个维度展开详细阐述。'
)

# 关键点1
add_heading('10.1  统一 AI 工作台与 Agent 自动路由', 2)

add_para('背景问题：', bold=True)
add_para(
    '在多业务场景的 AI 平台中，用户往往需要手动选择"用哪个智能体"，操作成本高且容易选错。'
    '例如用户想了解某学员的情况，但系统默认在"内容运营"智能体下，用户需要手动切换到"学员管理"智能体。'
    '这降低了平台的易用性，也增加了用户的学习成本。'
)

add_para('设计目标：', bold=True)
add_para('① 用户面对统一的 AI 对话框入口，不需要手动选择智能体类型。', bold=False)
add_para('② 系统自动识别用户意图，路由到最合适的智能体。', bold=False)
add_para('③ 支持用户显式指定 agent_type 覆盖自动路由结果。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    '设计了 AgentRouter（位于 ai-agent-service），其核心是关键词匹配的意图推断算法。'
    'AgentRouter 维护一个从 agent_type 到 Agent 实例的 Map，通过 inferAgentType() 方法根据用户输入文本中的关键词判断意图。'
    '判断优先级从具体到通用：先检查飞书文档创建意图（同时包含"飞书/文档"和"创建/生成"），'
    '再按 feishu → student → school → knowledge → content 的顺序匹配关键词，最后默认 fallback 到 ContentAgent。'
)

add_para('关键流程：', bold=True)
add_para('① 前端请求中 agent_type 默认为 "auto"。', bold=False)
add_para('② AgentRouter.inferAgentType() 执行关键词匹配。', bold=False)
add_para('③ 若用户显式指定了有效的 agent_type（非 auto），且不等于 "all"，则使用用户指定的 Agent。', bold=False)
add_para('④ 否则使用关键词推断的 Agent。', bold=False)
add_para('⑤ Agent 的 execute() 或 stream() 方法被调用，执行具体的智能体逻辑。', bold=False)

add_para('优点：', bold=True)
add_para('① 用户无需学习"切换智能体"的概念，直接输入问题即可。', bold=False)
add_para('② 关键词匹配规则透明可维护，新增领域（如"财务"、"法务"）只需添加新的关键词分支和对应的 Agent 实现。', bold=False)
add_para('③ 支持用户显式覆盖，兼具灵活性和易用性。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① 当前关键词匹配对复杂意图（如"帮我分析学员张三的 GPA 并推荐合适的院校"）无法同时触发多个 Agent。后续可引入意图分类模型（如基于 LLM 的 few-shot 分类）或将复杂任务分解为子任务。', bold=False)
add_para('② 关键词列表需要手动维护，新增业务领域时需要更新 inferAgentType() 方法。后续可考虑将关键词配置外置到配置文件或数据库。', bold=False)

# 关键点2
add_heading('10.2  LLM API 可替换封装', 2)

add_para('背景问题：', bold=True)
add_para(
    '不同大模型厂商的 API 接口存在差异（如 OpenAI 的 /v1/chat/completions、DeepSeek 的 /chat/completions、'
    '豆包/通义千问各自的原生接口）。如果业务代码直接调用某个厂商的 SDK，后续切换模型需要大量代码修改。'
    '此外，课程演示和本地调试时可能没有可用的 API Key，需要一个离线的 Mock 方案。'
)

add_para('设计目标：', bold=True)
add_para('① 业务代码不感知底层使用的大模型厂商。', bold=False)
add_para('② 支持流式（stream）和非流式（complete）两种调用方式。', bold=False)
add_para('③ 支持解析模型的 Thinking/Reasoning 字段。', bold=False)
add_para('④ 提供 Mock 实现用于离线 Demo。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    '定义 LLMClient 接口（位于 ai-agent-service/llm/），包含两个方法：complete() 用于非流式调用，'
    'stream() 用于流式调用。stream() 方法有两个重载：两个 Consumer 参数（onDelta + onReasoningDelta）'
    '和单 Consumer 参数（向后兼容）。'
)
add_para(
    '两个实现类：MockLLMClient 返回预设的模拟文本，包含模拟的 Thinking 和回答，不依赖任何外部 API Key；'
    'OpenAiCompatibleLLMClient 通过 HTTP 调用兼容 OpenAI 协议的大模型 API，支持通过 LlmProperties 配置 '
    'baseUrl、apiKey、model 等参数，可以对接 DeepSeek、OpenAI、豆包、通义千问以及任何兼容 OpenAI Chat Completions 协议的 API。'
)

add_para('优点：', bold=True)
add_para('① 切换大模型只需修改配置文件（改 baseUrl 和 model），无需修改任何业务代码。', bold=False)
add_para('② MockLLMClient 让开发者在没有 API Key 的情况下也能运行和测试完整流程。', bold=False)
add_para('③ stream() 方法的双 Consumer 设计使得 Reasoning/Thinking 和最终回答可以在前端分别展示。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① 当前仅适配 OpenAI 兼容协议。后续可扩展 Anthropic 协议适配器、Ollama 本地模型适配器。', bold=False)
add_para('② 流式解析逻辑（SSE 事件解析）与 LLMClient 耦合在 OpenAiCompatibleLLMClient 中。后续可抽象出 StreamParser 接口以支持不同的流式协议（如 Anthropic 的 Message Stream）。', bold=False)
add_para('③ 当前没有内置的 Token 计数和成本统计。后续可集成 Token 计数器并在 metadata 中记录每次调用的 Token 用量和费用。', bold=False)

# 关键点3
add_heading('10.3  SSE 流式回答、模型 Thinking 与工具 Trace 展示', 2)

add_para('背景问题：', bold=True)
add_para(
    '大模型的生成通常需要数秒到数十秒。如果使用传统的"请求-等待-完整返回"模式，用户可能需要等待 10-30 秒才能看到任何反馈，'
    '体验极差，且可能误以为系统卡死。此外，DeepSeek 等模型支持返回 thinking/reasoning_content 字段（模型的思考过程），'
    '如果不将这些信息展示给用户，就浪费了大模型的透明度能力。工具调用的输入输出如果不展示，用户也无法判断 AI 如何得出结论。'
)

add_para('设计目标：', bold=True)
add_para('① 流式返回 AI 回答，用户可以看到逐字生成过程。', bold=False)
add_para('② 分别展示模型的"思考过程"（Thinking）和"最终回答"（Answer）。', bold=False)
add_para('③ 展示 Agent 调用了哪些工具、每个工具的输入输出和耗时（Trace）。', bold=False)
add_para('④ 流式结束后所有信息持久化到数据库，支持历史回放。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    '后端 SSE 输出多种事件类型：'
)
add_para('• reasoning：模型返回的思考过程（reasoning_content 字段），前端追加到可折叠的"思考过程"区域。', bold=False)
add_para('• delta：模型返回的回答内容（content 字段），前端追加到主回答区。', bold=False)
add_para('• trace：工具调用结果，包含工具名称、输入摘要、输出摘要、耗时。前端展示在可折叠的"工具调用"区域。', bold=False)
add_para('• session：会话 ID，前端保存用于后续历史查询。', bold=False)
add_para('• done：流式结束信号，前端关闭 EventSource 连接。', bold=False)

add_para(
    '前端通过 EventSource API 连接后端的 /api/agent/stream 端点。收到不同事件类型后，更新 Vue 组件的响应式状态。'
    '思考过程和工具调用默认折叠，用户可点击展开查看详情。流式结束后，前端发送请求确认消息已完整接收，'
    '后端将完整的 Thinking、Trace 和 Answer 保存到 agent_message.metadata。'
)

add_para('优点：', bold=True)
add_para('① 流式输出显著提升用户感知速度——用户在看到第一个 token 时就知道系统在工作。', bold=False)
add_para('② Thinking 和 Trace 分开展示，让 AI 的推理过程透明化，用户和开发者都能理解 AI 决策链条。', bold=False)
add_para('③ 比 WebSocket 更轻量——不需要心跳维护、不需要全双工通信（AI 回答场景天然是单向推送）。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① SSE 仅支持服务端到客户端的单向推送，如果后续需要中断生成（如用户点击"停止生成"），需要通过另一个 REST 接口来实现。', bold=False)
add_para('② 浏览器对同一域名的 SSE 连接数有限制（通常 6 个），多标签页同时使用时可能阻塞。后续可考虑 HTTP/2 或迁移到 WebSocket。', bold=False)
add_para('③ 当前 metadata 使用了 MySQL 的 TEXT/JSON 字段而非专门的时序表，大量历史会话的 metadata 查询效率可能下降。后续可考虑将 Trace 数据迁移到独立的 agent_trace 表。', bold=False)

# 关键点4
add_heading('10.4  工具扩展/MCP/Skill 解耦设计', 2)

add_para('背景问题：', bold=True)
add_para(
    'AI 智能体平台需要不断集成新的工具能力：向量检索、飞书操作、联网搜索、时间工具、代码执行等。'
    '如果每个新工具都需要修改 Agent 核心逻辑，会导致主流程频繁变更、多人协作冲突频繁、代码耦合度上升。'
    '特别是在课程设计场景中，不同小组成员可能负责不同的工具能力开发，需要一个明确的扩展规范。'
)

add_para('设计目标：', bold=True)
add_para('① 新增工具能力时尽量少改动 Agent 主流程代码。', bold=False)
add_para('② 统一的工具描述接口，让 Agent 能自动感知可用工具。', bold=False)
add_para('③ 支持不同类型的扩展：MCP 工具（外部工具协议）、Skill（固定工作流）、Runtime 工具（时间/联网等内置能力）。', bold=False)
add_para('④ 小组成员可以在独立的包中开发工具能力，降低代码冲突。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    'AgentExtension 接口定义了统一的扩展契约：'
)
add_para('• name()：扩展名称（如 "向量检索"、"飞书MCP"）。', bold=False)
add_para('• type()：扩展类型（mcp / skill / runtime）。', bold=False)
add_para('• description()：扩展功能描述，用于 LLM 了解该工具的能力。', bold=False)
add_para('• supports(agentType)：该扩展支持哪些 Agent 类型，默认支持所有。', bold=False)
add_para('• runtimeContext(request)：返回该扩展在当前上下文中的运行时信息（如当前时间、向量检索结果等），被注入到 LLM 的 System Prompt 中。', bold=False)

add_para(
    '实现类包括：McpToolProvider（MCP 工具协议）、SkillProvider（SOP 工作流能力）、'
    'LarkCliMcpExtension（飞书 MCP 工具）、SemanticVectorSearchExtension（向量检索）、'
    'RuntimeToolExtensions（时间、联网等基础能力）。'
    '所有扩展通过 DefaultAgentExtensions 聚合，在 Agent 执行时按顺序收集 runtimeContext。'
)

add_para(
    '小组成员如需新增工具能力，只需在 contrib/capability 包下实现 AgentExtension 接口并注册为 Spring Bean，'
    '无需修改 AgentRouter、BaseAgent 或任何现有 Agent 的代码。工具能力稳定后，再迁移到正式的 agent.extension 包。'
)

add_para('优点：', bold=True)
add_para('① 工具能力和业务智能体解耦——ContentAgent 不依赖具体的飞书 SDK 或 Weaviate SDK，只依赖 AgentExtension 接口。', bold=False)
add_para('② 多人协作友好——小组成员各自在自己的 contrib 包中开发工具，主流程代码不变，减少合并冲突。', bold=False)
add_para('③ 工具能力可组合——一个 Agent 可以同时使用向量检索、飞书操作和联网搜索，Extensions 链按声明顺序依次提供上下文。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① 当前 runtimeContext() 是同步的，所有扩展按顺序执行。后续可考虑并行执行独立的扩展以提高响应速度。', bold=False)
add_para('② 扩展执行失败时缺乏细粒度的降级策略。后续可增加 fallbackContext() 方法，让每个扩展定义自己的降级返回内容。', bold=False)
add_para('③ 当前 MCP 工具通过 lark-cli 命令行桥接，进程启动开销较大。后续可替换为原生飞书 OpenAPI SDK 调用。', bold=False)

# 关键点5
add_heading('10.5  知识库向量检索方案', 2)

add_para('背景问题：', bold=True)
add_para(
    '传统的关键词搜索（MySQL LIKE）无法理解语义相似性。例如用户搜索"保研面试准备"，'
    '传统搜索只能匹配包含这些精确关键词的文档，无法匹配"推免复试技巧"或"夏令营面试经验"等语义相近但关键词不同的文档。'
    '此外，大模型自身有知识截止日期和幻觉问题，在回答关于特定院校项目或内部知识时可能生成不准确的内容。'
)

add_para('设计目标：', bold=True)
add_para('① 支持语义级相似度检索，不依赖精确关键词匹配。', bold=False)
add_para('② 检索结果带来源追溯（文档标题、飞书 Token、原文链接）。', bold=False)
add_para('③ 向量数据库不可用时能降级到 MySQL 文本搜索。', bold=False)
add_para('④ 检索结果作为上下文注入 LLM Prompt，提升回答的事实准确度。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    '知识库向量检索方案包含三个核心环节：'
)

add_para('（1）文档入库与向量化', bold=True)
add_para('• TextChunker：将知识文档按段落切分为 chunk（默认每 chunk 500-1000 字符，按段落边界切分以保持语义完整）。', bold=False)
add_para('• EmbeddingService：调用嵌入模型 API（兼容 OpenAI Embeddings 协议）为每个 chunk 生成向量。', bold=False)
add_para('• WeaviateClientService：通过 Weaviate GraphQL 接口将向量 + 元数据（title, chunkText, feishuToken, tags, mysqlId 等）写入 KnowledgeDoc class。', bold=False)

add_para('（2）语义检索', bold=True)
add_para('• VectorSearchToolService：接收查询文本 → EmbeddingService 生成查询向量 → Weaviate 近邻搜索 → 返回 TopK（默认 K=5）结果。', bold=False)
add_para('• 返回结果包含：source（文档标题）、chunkText（匹配分片原文）、feishuUrl（飞书链接）、tags（标签）、distance（相似度距离）。', bold=False)

add_para('（3）检索结果注入', bold=True)
add_para('• SemanticVectorSearchExtension 作为 AgentExtension 实现，在其 runtimeContext() 方法中调用 VectorSearchToolService。', bold=False)
add_para('• 检索结果作为结构化上下文注入到 LLM 的 System Prompt 中（例如："以下是知识库中与用户问题相关的资料：\n---\n{context}\n---\n请基于以上资料回答用户问题。如果资料中没有相关信息，请如实说明。"）。', bold=False)
add_para('• 当 Weaviate 不可用或返回空结果时，降级为 MySQL 的 LIKE 搜索。', bold=False)

add_para('优点：', bold=True)
add_para('① RAG（检索增强生成）模式显著减少了大模型的幻觉——回答有了知识库的事实依据。', bold=False)
add_para('② 来源可追溯——每段回答都可以对应到具体的知识文档和原始出处。', bold=False)
add_para('③ Weaviate 和 MySQL 的分工清晰：Weaviate 做语义检索，MySQL 做结构化查询和管理。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① 当前 Embedding 模型由外部 API 提供，后续可考虑使用本地 Embedding 模型（如 BGE、M3E）以降低延迟和成本。', bold=False)
add_para('② TextChunker 使用固定长度切分，可能切断语义段落。后续可引入语义分块（基于句子边界或 LLM 判断的语义边界）。', bold=False)
add_para('③ 当知识文档更新或删除时，Weaviate 中的对应向量需要同步更新。当前依赖手动触发或定时同步，后续可实现基于事件驱动的实时同步。', bold=False)

# 关键点6
add_heading('10.6  飞书生态接入方案', 2)

add_para('背景问题：', bold=True)
add_para(
    '飞书提供文档、多维表格、任务、群机器人等多种能力，但其 API 体系复杂（需要 tenant_access_token、app_access_token、'
    '不同 API 有不同的权限范围）。直接在业务模块中编写飞书 API 调用代码会导致：① 代码耦合度高；'
    '② 更换飞书 SDK 或 API 版本时需要大量修改；③ 课程 Demo 阶段可能没有完整的飞书应用权限。'
)

add_para('设计目标：', bold=True)
add_para('① 统一封装飞书能力，业务模块不直接依赖飞书 SDK。', bold=False)
add_para('② Demo 阶段支持 Mock 飞书数据，不依赖真实飞书权限。', bold=False)
add_para('③ 飞书操作结果进入工具 Trace，便于追踪和审计。', bold=False)
add_para('④ 文档共享文件夹通过 folderToken 访问，创建文档时通过 parentToken 指定目标位置。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    '飞书能力封装分为三个层次：'
)
add_para('① feishu-service 模块：提供 FeishuController（REST 接口）和 FeishuMockClient（Mock 实现），封装飞书文档、多维表格、任务、机器人的基本操作。', bold=False)
add_para('② LarkCliMcpExtension：作为 AgentExtension 的 MCP 实现，通过 lark-cli 命令行工具桥接飞书 API。Agent 通过 AgentExtension 接口调用飞书能力，不感知底层是 lark-cli 还是飞书 SDK。', bold=False)
add_para('③ FeishuAgent：专门处理飞书相关任务的智能体，如"在飞书知识库文件夹中创建一篇关于保研面试的文档"。', bold=False)

add_para('关键流程（飞书文档创建）：', bold=True)
add_para('① 用户在 AI 工作台中输入："在飞书保研知识库文件夹里创建一篇关于北京大学计算机夏令营的文档"。', bold=False)
add_para('② AgentRouter 识别为飞书文档创建请求，路由到 FeishuAgent。', bold=False)
add_para('③ FeishuAgent 通过 LLM 生成文档标题和正文内容。', bold=False)
add_para('④ LarkCliMcpExtension 调用 lark-cli 创建飞书文档，指定 parentToken 为目标文件夹。', bold=False)
add_para('⑤ 创建成功后记录飞书文档 URL 和 Token 到 feishu_sync_record 表。', bold=False)
add_para('⑥ 工具调用 Trace 记录创建结果（成功/失败、文档 URL、耗时）。', bold=False)

add_para('优点：', bold=True)
add_para('① 解耦：业务 Agent 不依赖飞书 SDK，后续更换为原生飞书 OpenAPI 时只需替换 LarkCliMcpExtension 的实现。', bold=False)
add_para('② Demo 友好：无飞书权限时 FeishuMockClient 返回模拟数据，不影响系统演示。', bold=False)
add_para('③ 可追溯：所有飞书操作的结果记录在工具 Trace 和 sync_record 中。', bold=False)
add_para('④ 权限透明：飞书 Token 过期或权限不足时，系统明确提示用户，不静默失败。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① lark-cli 作为外部进程调用，存在进程启动开销（约 200-500ms）。后续可替换为原生飞书 OpenAPI Java SDK 或 feign 客户端。', bold=False)
add_para('② 当前飞书操作结果不自动通知用户（如飞书文档创建成功后只在前端 Trace 中展示）。后续可集成飞书群机器人消息推送。', bold=False)

# 关键点7
add_heading('10.7  内容运营数据沉淀设计', 2)

add_para('背景问题：', bold=True)
add_para(
    '内容运营人员在使用 AI 生成文案时，通常面临以下问题：生成的文案可能有多版修改、需要排期发布、'
    '发布后需要追踪效果、好的选题需要沉淀为资产。如果系统只提供"一次性生成"，'
    '那么运营人员每次生成都需要重新描述需求，历史好文案无法复用，选题灵感无法积累。'
)

add_para('设计目标：', bold=True)
add_para('① 内容生成不是一次性输出，而是一个资产沉淀过程。', bold=False)
add_para('② 选题、文案、图片、日历、评分、发布指标形成完整的内容资产链条。', bold=False)
add_para('③ 支持跨渠道（小红书、朋友圈、公众号、飞书文档）的内容管理。', bold=False)
add_para('④ 支持文案评分和效果追踪，为后续 AI 优化提供数据基础。', bold=False)

add_para('方案设计：', bold=True)
add_para(
    'content-service 模块围绕内容资产全生命周期设计了以下数据实体和流程：'
)
add_para('① 主题库（content_topic）：保存选题，标注平台、类型、风格、热度分、收藏状态。好选题不会丢失，可反复使用。', bold=False)
add_para('② 文案库（content_copy）：保存多个版本文案（初稿、优化稿、发布稿），支持 1-5 星评分和使用状态追踪。', bold=False)
add_para('③ 内容日历（content_calendar）：关联选题，记录计划发布日期和渠道，形成内容排期视图。', bold=False)
add_para('④ 配图管理（content_copy_image）：单独管理文案配图或建议图片。', bold=False)
add_para('⑤ 生成流水（content_generation_record）：保存每次生成的输入摘要、输出摘要、使用的 SOP 类型，用于后续分析哪些 SOP 效果好。', bold=False)
add_para('⑥ 发布指标（content_publish_metric）：记录发布后的实际效果数据（阅读量、互动量、转化量），为 AI 文案优化提供反馈闭环。', bold=False)

add_para('优点：', bold=True)
add_para('① 内容资产可积累：随着使用，内容主题库和文案库不断丰富，AI 生成的新文案可以借鉴历史好文案。', bold=False)
add_para('② 反馈闭环：发布指标数据可以用于评估 AI 文案质量，形成"生成-发布-评估-优化"的闭环。', bold=False)
add_para('③ 跨平台统一管理：同一选题可以在不同平台生成不同风格的文案，统一在内容日历中排期。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① 当前发布指标需要手动录入，后续可对接各平台（小红书、公众号）的开放 API 自动获取数据。', bold=False)
add_para('② 内容 SOP 目前为固定流程，后续可支持运营人员自定义 SOP 模板。', bold=False)
add_para('③ 内容资产目前仅本地存储，后续可同步到飞书多维表格，利用飞书的协作和分享能力。', bold=False)

# 关键点8
add_heading('10.8  安全配置与敏感信息管理', 2)

add_para('背景问题：', bold=True)
add_para(
    '在开发过程中，API Key、数据库密码等敏感信息极易通过配置文件泄露到 Git 仓库。'
    '一旦推送到公开仓库，攻击者可以利用泄露的 Key 调用大模型 API 产生费用，或访问数据库造成数据泄露。'
    '此外，前端如果持有 LLM API Key，任何人通过浏览器开发者工具即可获取，造成严重安全隐患。'
)

add_para('设计目标：', bold=True)
add_para('① API Key、密码等敏感信息不进入 Git 版本控制。', bold=False)
add_para('② 前端不持有 LLM Key 或飞书 Key。', bold=False)
add_para('③ 后端统一代理所有外部 API 调用。', bold=False)
add_para('④ 工具调用有安全边界，不暴露危险操作给 LLM。', bold=False)

add_para('方案设计：', bold=True)
add_para('① application-local.yml 文件存储本地敏感配置（数据库密码、LLM API Key、飞书 App Secret），该文件已加入 .gitignore，不进入 Git。', bold=False)
add_para('② 提供 application-local.yml.example 模板文件，列出需要配置的 Key（但不包含实际值），供新开发者参考。', bold=False)
add_para('③ 所有外部 API 调用（LLM、Embedding、飞书）均经过后端 Controller → Service → Adapter 链路，前端只与自己的后端通信，不直接访问第三方 API。', bold=False)
add_para('④ MockLLMClient 作为无 Key 状态下的降级方案：当 application-local.yml 中未配置 LLM API Key 时，系统自动使用 MockLLMClient 返回模拟数据，保证系统可运行。', bold=False)
add_para('⑤ 工具调用白名单机制：Agent 只能通过已注册的 AgentExtension 调用工具，不存在"执行任意命令"的扩展。', bold=False)

add_para('优点：', bold=True)
add_para('① 敏感信息不进入 Git，从根本上避免了泄露风险。', bold=False)
add_para('② 前端零敏感信息，即使用户打开浏览器 DevTools 也无法获取 API Key。', bold=False)
add_para('③ Mock 模式让无 API Key 的开发者和评审老师也能运行系统。', bold=False)

add_para('局限与后续优化：', bold=True)
add_para('① application-local.yml 依赖于开发者自觉不提交。后续可引入 git-secrets 或 .gitallowed 机制做自动化检查。', bold=False)
add_para('② 生产环境建议使用 Vault（HashiCorp）或云厂商的密钥管理服务（如阿里云 KMS）管理敏感配置，而非本地文件。', bold=False)
add_para('③ 对于需要用户自行提供 API Key 的场景（如用户使用自己的 DeepSeek Key），后续可提供"用户级 API Key 管理"功能，Key 加密存储在数据库中。', bold=False)

add_heading('10.9  桌面客户端高保真迁移方案', 2)

add_para('背景问题：', bold=True)
add_para(
    '原 FlowMind Agent 仅提供了 Web 前端（Vue 3 SPA），但在课程演示、离线环境或用户偏好桌面应用场景下，'
    'Web 端存在以下局限：① 需要先启动 Vite Dev Server 或构建静态文件部署到 Nginx；'
    '② 浏览器环境无法直接访问本地文件系统或系统通知；③ 课程演示时网络不稳定可能导致前端资源加载失败。'
)

add_para('设计目标：', bold=True)
add_para('① 用 Python/PySide6 高保真复刻 Web 端全部 10 个业务页面，保持页面结构和交互流程一致。')
add_para('② 复用现有 Java 后端，桌面端只重写客户端界面，不重复开发业务逻辑。')
add_para('③ 支持离线 Demo 模式，无后端无网络时仍可完整演示所有页面。')
add_para('④ 提供自动化自检脚本，验证 API、离线数据和 GUI 页面是否正常工作。')

add_para('方案设计：', bold=True)
add_para(
    '桌面端采用"保留后端、重写客户端"策略，核心设计决策如下：'
)
add_para('① API 客户端层（api.py）：统一封装所有后端 REST 和 SSE 调用，同时内置 offline://demo 模式的完整 Mock 数据（fallback_data.py），当后端不可用时自动切换为离线模式。'
         '所有 Mock 数据覆盖登录、角色权限、流式对话、内容运营、知识库、学员、院校、飞书状态等全部业务模块。')
add_para('② 页面视图层（views.py）：每个页面实现为一个独立的 Python 类，包含 load() 数据加载、render() 页面渲染和用户交互处理。'
         '页面间通过 QStackedWidget 切换，侧边栏导航与 Web 端路由结构一致。')
add_para('③ 通用组件层（widgets.py）：将 Web 端的 Element Plus 组件等价映射为 PySide6 自定义组件：'
         'Badge → 彩色标签、StarRating → 星级评分、MarkdownPanel → Markdown 渲染、'
         'TraceListPanel → 工具调用卡片列表、CalendarGrid → 月历网格、ChartCard → 自绘图表（QPainter）、'
         'DataTable → 可交互数据表格、StreamWorker → SSE 流式线程。')
add_para('④ 样式系统（styles.py）：通过 QSS（Qt Style Sheets）定义全应用统一的设计语言，'
         '包括颜色系统（PRIMARY #5b6cff、SUCCESS #19b37b、WARNING #f59e0b、DANGER #ef4444）、'
         '字体规范（Microsoft YaHei）、圆角/边框/间距规范、按钮/标签/卡片/进度条的视觉变体。')
add_para('⑤ 双渲染器设计：PySide6 作为主渲染器（run.py），提供接近原生的现代 GUI 体验；'
         'Tkinter（tk_app.py）作为兜底渲染器，在 PySide6 因 Qt DLL 兼容性问题无法启动时仍可运行核心功能。')

add_para('优点：', bold=True)
add_para('① 一次开发，多端共享：后端 API 同时服务 Web、桌面和移动三个客户端，业务逻辑零重复。')
add_para('② 离线友好：内置 fallback_data.py 覆盖全部 Mock 数据，评审老师无需启动后端即可查看完整界面。')
add_para('③ 自检验证：smoke_test.py + pyside_smoke_test.py + check_desktop.py 三层自检确保提交质量。')
add_para('④ 渐进增强：从 Tkinter 兜底 → PySide6 高保真，不依赖单一 GUI 框架。')

add_para('局限与后续优化：', bold=True)
add_para('① PySide6 6.11+ 曾出现 Qt DLL 导入失败，当前固定 6.7.3 版本以保持兼容。后续可升级到更新的稳定版本。')
add_para('② 自绘图表（ChartCanvas）不支持交互式缩放和 tooltip，后续可集成 pyqtgraph 或 Qt Charts 模块。')
add_para('③ 桌面端目前不支持 Web 端的 Service Worker 离线缓存和 PWA 安装能力。')

add_heading('10.10  移动端 Android 客户端', 2)

add_para('背景问题：', bold=True)
add_para(
    '内容运营人员和教育咨询老师经常在外出或移动场景下需要快速查看知识库、回答学员问题、'
    '检查内容日历。Web 端在手机上体验不佳（未做移动端响应式适配），需要一个原生 Android 客户端。'
)

add_para('设计目标：', bold=True)
add_para('① 提供原生 Android 应用，支持在手机上使用 AI 工作台核心功能。')
add_para('② 与 Web 端共享同一套后端 API 和用户认证体系，数据互通。')
add_para('③ 支持 SSE 流式对话，移动端也能体验流式输出。')

add_para('方案设计：', bold=True)
add_para(
    '移动端采用原生 Android (Java) 开发，核心组件包括：'
)
add_para('① ApiClient.java（~198 行）：统一的 HTTP 客户端，封装所有后端 REST 接口调用和 SSE 流式数据解析。'
         '采用异步线程 + Handler 机制，将流式 token 实时更新到 UI。')
add_para('② MainActivity.java（~1,427 行）：主界面实现，包含对话列表、消息气泡、输入框和流式文本展示。'
         '支持多 Agent 类型切换（FlowMindAgent / ContentAgent / KnowledgeAgent / StudentAgent / SchoolAgent / FeishuAgent）。')
add_para('③ 编译产物为 Debug APK（FlowMindAgent-mobile-debug.apk，~67KB），可直接安装到 Android 手机进行功能验证。')

add_para('优点：', bold=True)
add_para('① 原生体验：利用 Android 原生 UI 组件，滑动流畅、启动快。')
add_para('② 后端复用：与 Web 端和桌面端共享完全相同的后端 API，无需为移动端单独开发服务端。')
add_para('③ 鉴权一致：使用与 Web 端相同的 Token 认证机制，用户账号在三个端通用。')

add_para('局限与后续优化：', bold=True)
add_para('① 当前为 Debug 构建，未做 ProGuard 混淆和 Release 签名。后续需配置正式签名和构建流程。')
add_para('② UI 较简单，未使用 Jetpack Compose 或 Material Design 3。后续可升级 UI 框架提升视觉效果。')
add_para('③ 目前仅实现核心对话和列表功能，内容运营、知识库管理等高级页面的移动端适配仍需完善。')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第十一章 系统附加说明
# ═══════════════════════════════════════════════
add_heading('11  系统附加说明', 1)

add_heading('11.1  运行环境', 2)

add_table(
    ['环境项', '要求', '说明'],
    [
        ['操作系统', 'Windows 11（开发环境）/ macOS / Linux', 'Docker Desktop 用于运行 MySQL 和 Weaviate'],
        ['JDK', '17 或以上', '建议 Azul Zulu 17 或 OpenJDK 17'],
        ['Node.js', '18 或以上', '前端开发和构建'],
        ['Maven', '3.8+（项目内置 Maven Wrapper）', '无需手动安装 Maven，使用 mvnw 脚本'],
        ['Docker Desktop', '最新稳定版', '运行 MySQL 9.x 和 Weaviate 容器'],
        ['MySQL', '9.x（Docker 容器）', '端口 3306，数据库名 FlowMind'],
        ['Weaviate', '最新稳定版（Docker 容器）', '端口 8081，向量维度和 Embedding 模型一致'],
        ['lark-cli', '最新版（可选）', '飞书命令行工具，用于飞书相关功能。无 lark-cli 时可使用 Mock 模式'],
        ['浏览器', 'Chrome 90+ / Edge 90+', '需要支持 SSE（EventSource）和 ES Module'],
    ]
)

add_heading('11.2  开发环境', 2)

add_table(
    ['工具', '用途'],
    [
        ['IntelliJ IDEA / VS Code', '后端 Java 开发 / 前端 Vue 开发'],
        ['Git', '版本控制'],
        ['PowerShell / Bash', '命令行操作'],
        ['npm', '前端包管理'],
        ['Maven Wrapper (mvnw)', '后端构建和依赖管理'],
        ['Docker', '数据库和向量数据库容器管理'],
        ['Postman / Swagger UI', 'API 调试'],
        ['Navicat / DBeaver', '数据库管理（可选）'],
    ]
)

add_heading('11.3  安装方法', 2)

add_para('详细安装步骤：', bold=True)
add_para('1. 克隆仓库：git clone <repo-url> && cd FlowMind-Agent', bold=False)
add_para('2. 启动 MySQL：执行 docker run -d --name flowmind-mysql -p 3306:3306 -e MYSQL_ROOT_PASSWORD=<your-password> mysql:9，然后执行 backend/sql/schema.sql 创建表结构，再执行 backend/sql/mock-data.sql 导入示例数据。', bold=False)
add_para('3. 启动 Weaviate：执行 docker run -d --name flowmind-weaviate -p 8081:8080 -e AUTHENTICATION_ANONYMOUS_ACCESS_ENABLED=true semitechnologies/weaviate:latest', bold=False)
add_para('4. 创建 backend/application-local.yml，参考 application-local.yml.example 填入本地配置（数据库密码、LLM API Key 等）。如不填 Key，系统自动使用 Mock 模式。', bold=False)
add_para('5. 启动后端：cd backend && ./mvnw spring-boot:run（Windows 下使用 mvnw.cmd spring-boot:run），确认控制台输出 "FlowMind Application started"。', bold=False)
add_para('6. 启动前端：cd frontend && npm install && npm run dev，确认控制台输出 Vite 本地地址（通常为 http://localhost:5173）。', bold=False)
add_para('7. 浏览器访问 http://localhost:5173，使用 Demo 账号登录（账号密码见 application-local.yml 或 Mock 数据）。', bold=False)

add_heading('11.4  卸载方法', 2)

add_para('卸载步骤：', bold=True)
add_para('1. 停止前端服务（Ctrl+C 停止 Vite Dev Server）。', bold=False)
add_para('2. 停止后端服务（Ctrl+C 停止 Spring Boot）。', bold=False)
add_para('3. 停止 Docker 容器：docker stop flowmind-mysql flowmind-weaviate', bold=False)
add_para('4. 删除本地配置文件（可选）：rm backend/application-local.yml', bold=False)
add_para('5. 删除 Docker 数据卷（可选，会清空所有数据）：docker rm flowmind-mysql flowmind-weaviate && docker volume prune', bold=False)
add_para('⚠ 注意：卸载前应备份 MySQL 数据库（使用 mysqldump 或 Navicat 导出）和飞书中的重要资料。Docker 数据卷删除后数据不可恢复。', bold=False)

add_heading('11.5  系统规模', 2)

add_para('（以下数据来自 git ls-files 统计，排除 node_modules、target、.venv、dist、.git 等非源码目录。）', bold=True)

add_table(
    ['类别', '代码量', '统计方式'],
    [
        ['前端代码量（Vue/TS/JS/CSS）', '~9,045 行', 'Vue 4,732 + TypeScript 4,311 + CSS/HTML'],
        ['服务端代码量（Java）', '~17,497 行', '103 个 Java 源文件，涵盖 12 个 Maven 模块'],
        ['桌面客户端代码量（Python）', '~7,220 行', 'views.py + widgets.py + api.py + styles.py + fallback_data.py + tk_app.py'],
        ['移动端代码量（Android Java）', '~1,625 行', 'MainActivity.java + ApiClient.java'],
        ['数据库脚本（SQL）', '~316 行', 'schema.sql + mock-data.sql'],
        ['文档与 Markdown', '~9,429 行', '所有 .md 文件'],
        ['配置与脚本（XML/YAML/Gradle/PS1/Sh）', '~1,239 行', 'pom.xml、application.yml、build.gradle、.ps1 等'],
        ['整个系统总代码量', '~49,879 行', '使用 git ls-files | xargs wc -l 在根目录统计全部源文件（290 个文件）'],
    ]
)

add_para(
    '推荐使用 cloc（Count Lines of Code）工具进行代码量统计。安装方式：npm install -g cloc 或 choco install cloc（Windows）。'
    '统计时排除 node_modules、target、dist、.m2repo、.venv、__pycache__、.git 等非源码目录。'
    '也可以使用 PowerShell 命令：git ls-files | Where-Object {$_ -match "\\.(java|vue|ts|py|js|css|html|sql|xml|yml|yaml|md)$"} | ForEach-Object {Get-Content $_ | Measure-Object -Line} | Measure-Object -Property Lines -Sum'
)

add_heading('11.6  待补充项', 2)

add_para('以下内容需要在系统最终提交前补充（请勿编造）：')
add_table(
    ['序号', '待补充内容', '说明'],
    [
        ['1', '最终运行截图', '包括 AI 工作台、内容管理、学员管理、数据分析、飞书操作等各页面的实际运行截图'],
        ['2', '最终代码量统计', '使用 cloc 工具统计后填入 11.5 节'],
        ['3', '最终成员分工', '列出每位小组成员负责的模块和贡献'],
        ['4', '最终部署截图', 'Docker 容器运行状态、后端启动日志、前端页面截图'],
        ['5', '最终测试结果', '功能测试用例和测试结果（或 Postman 接口测试截图）'],
        ['6', '最终飞书授权截图', '飞书应用管理后台的权限配置截图（如有）'],
        ['7', '最终数据库表结构截图', 'Navicat/DBeaver 中的 ER 图或表列表截图'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 第十二章 总结
# ═══════════════════════════════════════════════
add_heading('12  总结', 1)

add_para(
    '《FlowMind Agent：基于飞书生态的 AI 内容运营与知识管理智能体平台》系统设计文档全面阐述了系统的技术方案。'
    '系统采用前端 Vue 3 + 后端 Spring Boot 多模块的前后端分离架构，在 Demo 阶段通过 app-service 聚合启动的同时，'
    '保留了清晰的微服务边界和后续拆分路径。'
)

add_para(
    '系统设计的核心亮点包括：'
)
add_para('① 统一 AI 工作台与 Agent 自动路由机制，降低用户操作成本。', bold=False)
add_para('② LLM API 可替换封装，实现大模型厂商无依赖。', bold=False)
add_para('③ SSE 流式输出 + Thinking/Trace 分离展示，提升 AI 交互透明度和用户体验。', bold=False)
add_para('④ AgentExtension 工具扩展机制，实现工具能力和业务智能体的解耦。', bold=False)
add_para('⑤ MySQL + Weaviate 混合存储架构，兼顾结构化数据管理能力和语义检索能力。', bold=False)
add_para('⑥ Port/Adapter 架构模式在 content-service 中的应用，实现业务逻辑与外部实现的清晰分离。', bold=False)
add_para('⑦ 飞书生态的深度集成设计，覆盖文档、多维表格、任务、机器人等多种能力。', bold=False)
add_para('⑧ 内容运营数据的全生命周期管理，从选题、生成、排期到发布效果追踪形成完整闭环。', bold=False)
add_para('⑨ 敏感信息的分层管理机制，从 .gitignore 到应用配置到环境变量，逐级加强安全性。', bold=False)
add_para('⑩ contrib/capability 扩展缓冲区设计，支持小组成员独立开发和渐进式能力沉淀。', bold=False)
add_para('⑪ 桌面客户端（Python/PySide6）高保真复刻 Web 端全部 10 个页面，支持离线 Demo 和双渲染器兜底。', bold=False)
add_para('⑫ Android 移动客户端，与 Web/桌面端共享同一后端 API 和用户认证体系。', bold=False)

add_para(
    '作为课程设计 Demo，本系统在功能完整度和架构前瞻性之间取得了良好的平衡。'
    '系统在保证课程演示稳定的同时，为后续的生产化部署和功能扩展预留了明确的升级路径。'
    '文档中标注为"预留"和"待补充"的部分，将在后续版本的迭代中逐步完善。'
)

# ═══════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'docs', 'submission', '系统设计文档.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
