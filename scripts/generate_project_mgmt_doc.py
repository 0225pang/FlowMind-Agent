"""
生成《项目管理文档》Word 文档
FlowMind Agent：基于飞书生态的 AI 内容运营与知识管理智能体平台
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1: hs.font.size = Pt(16)
    elif level == 2: hs.font.size = Pt(14)
    else: hs.font.size = Pt(13)

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    if bold: run.font.bold = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════════════════ 封面 ═══════════════════════════
for _ in range(6): doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run('项目管理文档'); run.font.size = Pt(26); run.font.bold = True
run.font.name = '黑体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sp.add_run('FlowMind Agent\n基于飞书生态的 AI 内容运营与知识管理智能体平台')
run.font.size = Pt(18); run.font.name = '黑体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4): doc.add_paragraph()
ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run('版本：v1.0    日期：2026年6月    状态：课程设计 Demo').font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════ 文档修订记录 ═══════════════════════════
add_heading('文档修订记录', 1)
add_table(
    ['版本', '日期', '修订人', '修订内容'],
    [
        ['v0.1', '2026-06-01', '小组协作', '创建文档框架，填写项目基本信息和初步分工'],
        ['v0.5', '2026-06-10', '小组协作', '补充开发计划、协作规范、目录结构说明'],
        ['v0.8', '2026-06-15', '小组协作', '补充风险管理、质量保证、沟通记录'],
        ['v1.0', '2026-06-19', '小组协作', '定稿：补充最终成员分工、待补充项清单、交付物清单'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 目录 ═══════════════════════════
add_heading('目  录', 1)
add_para('（在 Word 中插入自动目录：引用 → 目录 → 自动目录）')
doc.add_page_break()

# ═══════════════════════════ 第一章 项目基本信息 ═══════════════════════════
add_heading('1  项目基本信息', 1)

add_heading('1.1  项目概述', 2)
add_para('项目名称：FlowMind Agent：基于飞书生态的 AI 内容运营与知识管理智能体平台', bold=True)
add_para('项目简称：FlowMind Agent', bold=True)
add_para('项目类型：软件工程课程结课设计（面向对象技术与方法）', bold=True)
add_para('开发周期：2026年5月20日 - 2026年6月20日（共31天）', bold=True)
add_para('项目仓库：GitHub（https://github.com/0225pang/FlowMind-Agent）', bold=True)

add_heading('1.2  项目定位', 2)
add_para(
    'FlowMind Agent 是一个面向内容运营、教育服务、个人 IP 运营和小型团队协作场景的 AI 智能体平台。'
    '系统以网页端 AI 对话工作台为统一入口，结合飞书生态中的文档、多维表格、任务协作、群机器人等能力，'
    '配合本地业务数据库、向量数据库和大模型 API，实现内容选题生成、资料库整理、知识问答、学员信息管理、'
    '院校项目情报管理、数据可视化和任务协同。'
)

add_heading('1.3  项目目标', 2)
add_para('本项目作为软件工程课程结课设计 Demo，核心目标如下：')
add_para('① 构建一个前后端分离的、具有明确业务价值的 AI 智能体平台，而非简单的聊天机器人。')
add_para('② 在 Demo 阶段完成 AI 工作台、内容运营、知识管理、学员管理、院校情报、数据分析、飞书同步七大核心模块的开发。')
add_para('③ 通过 Agent Router 自动路由、LLMClient 抽象、AgentExtension 工具扩展等机制，展示良好的软件架构设计能力。')
add_para('④ 输出完整的软件工程文档体系（需求文档、设计文档、项目管理文档）。')
add_para('⑤ 合理运用 Git 团队协作、模块化开发、分支管理等软件工程实践。')

add_heading('1.4  开发环境与工具', 2)
add_table(
    ['类别', '工具/环境', '版本/说明'],
    [
        ['操作系统', 'Windows 11', '主要开发环境'],
        ['IDE（后端）', 'IntelliJ IDEA', 'Java 开发与调试'],
        ['IDE（前端）', 'VS Code', 'Vue 3 / TypeScript 开发'],
        ['JDK', 'OpenJDK 17', 'LTS 版本'],
        ['Node.js', '18+', '前端构建与包管理'],
        ['构建工具', 'Maven Wrapper (mvnw)', '后端依赖管理与构建'],
        ['数据库', 'MySQL 9.x (Docker)', '结构化业务数据存储'],
        ['向量数据库', 'Weaviate (Docker)', '知识库语义检索'],
        ['版本控制', 'Git + GitHub', '代码版本管理与团队协作'],
        ['项目管理', 'GitHub Issues / Projects', '任务追踪（推荐）'],
        ['文档', 'Microsoft Word / Markdown', '工程文档撰写'],
        ['API 调试', 'Postman / Swagger UI', '接口测试'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 第二章 项目团队成员与分工 ═══════════════════════════
add_heading('2  项目团队成员与分工', 1)

add_heading('2.1  团队组织架构', 2)
add_para('项目团队采用以组长为核心的协作模式，组长负责整体架构设计、主分支管理和代码审查，'
         '各成员根据分工独立负责特定模块的开发。组织架构如下：')

add_table(
    ['角色', '职责描述', '对应人员（待填写）'],
    [
        ['组长 / 架构师', '整体系统架构设计、核心模块开发、主分支管理、代码审查与合并、技术难点攻关、文档审核', '待补充'],
        ['后端开发（AI 智能体）', 'ai-agent-service 模块开发：Agent Router、LLMClient、AgentExtension、SSE 流式输出、会话管理', '待补充'],
        ['后端开发（内容运营）', 'content-service 模块开发：内容 SOP、主题库/文案库/日历、Port/Adapter 架构实现', '待补充'],
        ['后端开发（知识库 + 飞书）', 'knowledge-service 和 feishu-service 开发：向量检索、Weaviate 集成、飞书 API 桥接', '待补充'],
        ['后端开发（业务服务）', 'user-service、student-service、school-service、analytics-service 开发：用户认证、学员管理、院校情报、数据分析', '待补充'],
        ['前端开发（AI 工作台 + 核心页面）', 'AI 工作台界面、SSE 流式展示、Thinking/Trace 面板、主布局、路由守卫', '待补充'],
        ['前端开发（业务页面）', '内容运营、知识库、学员管理、院校情报、数据分析、飞书同步、系统设置页面', '待补充'],
        ['桌面客户端开发', 'desktop_fronted 桌面客户端（Python/Tkinter 或 PySide）开发', '待补充'],
        ['文档与测试', '系统需求文档、设计文档、项目管理文档撰写、功能测试与截图、演示脚本准备', '待补充'],
    ]
)

add_heading('2.2  成员能力矩阵', 2)
add_para('以下为团队成员技术能力分布（建议填写），用于明确各成员的技能覆盖情况：')

add_table(
    ['技能领域', '成员A', '成员B', '成员C', '成员D', '成员E（如有）'],
    [
        ['Java / Spring Boot', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['Vue 3 / TypeScript', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['Python / 桌面开发', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['MySQL / 数据库设计', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['Docker / 运维', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['LLM / AI 应用', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['Git / 协作', '待补充', '待补充', '待补充', '待补充', '待补充'],
        ['技术文档写作', '待补充', '待补充', '待补充', '待补充', '待补充'],
    ]
)

add_heading('2.3  工作量分配', 2)
add_para('以下为各模块的预估工作量和负责人（待最终填写）：')

add_table(
    ['模块/子系统', '预估工作量（人天）', '负责人', '状态'],
    [
        ['系统架构设计与数据库设计', '待补充', '待补充', '已完成'],
        ['ai-agent-service（Agent Router + LLMClient + Extension + SSE）', '待补充', '待补充', '已完成'],
        ['content-service（内容 SOP + 主题库/文案库/日历）', '待补充', '待补充', '已完成'],
        ['knowledge-service（Weaviate 集成 + 向量检索 + 文档管理）', '待补充', '待补充', '已完成'],
        ['feishu-service（飞书 API 桥接 + Mock + 同步日志）', '待补充', '待补充', '已完成'],
        ['user-service（认证 + 角色权限 + RBAC 预留）', '待补充', '待补充', '已完成'],
        ['student-service + school-service + analytics-service', '待补充', '待补充', '已完成'],
        ['前端 AI 工作台（SSE + Thinking/Trace + 会话）', '待补充', '待补充', '已完成'],
        ['前端业务页面（内容/知识库/学员/院校/数据分析/飞书/设置）', '待补充', '待补充', '已完成'],
        ['desktop_fronted 桌面客户端', '待补充', '待补充', '已完成'],
        ['文档系统（需求文档 + 设计文档 + 项目管理文档）', '待补充', '待补充', '完成中'],
        ['集成测试与截图', '待补充', '待补充', '进行中'],
        ['合计', '待统计', '—', '—'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 第三章 组织方式与协作规范 ═══════════════════════════
add_heading('3  组织方式与协作规范', 1)

add_heading('3.1  代码仓库与分支管理', 2)
add_para('项目代码托管于 GitHub：https://github.com/0225pang/FlowMind-Agent', bold=True)

add_heading('3.1.1  分支策略', 3)
add_para('项目采用简化版 Git Flow 分支策略：')
add_para('• main 分支：主分支，始终保持可运行状态。所有功能开发完成后通过 Pull Request 或直接合并入 main。')
add_para('• 功能分支（feature/xxx）：各成员从 main 拉出功能分支进行独立开发，'
         '命名格式为 feature/<模块名>-<简要描述>，例如 feature/content-sop、feature/knowledge-vector-search。')
add_para('• 修复分支（fix/xxx）：用于修复 main 分支上的 Bug，命名格式为 fix/<问题描述>。')
add_para('• 文档分支（docs/xxx）：用于文档编写和更新，命名格式为 docs/<文档类型>。')

add_heading('3.1.2  协作流程', 3)
add_para('① 组长创建 GitHub 仓库，初始化项目结构，搭建前端和后端基础框架，推送到 main 分支。')
add_para('② 各成员从 main 拉出功能分支：git checkout -b feature/<模块名>')
add_para('③ 成员在功能分支上独立开发，定期提交（小步提交，清晰的 commit message）。')
add_para('④ 功能开发完成后，成员将功能分支推送到远程仓库，由组长进行代码审查。')
add_para('⑤ 组长确认无误后，将功能分支合并到 main 分支。如有冲突，由组长协调解决。')
add_para('⑥ 合并后删除远程功能分支（本地可保留），成员切回 main 并拉取最新代码后开始下一轮开发。')

add_heading('3.1.3  Commit 规范', 3)
add_para('为保持提交历史的清晰可读，约定以下 Commit Message 格式：')
add_para('格式：<type>: <简短描述>', bold=True)
add_table(
    ['Type', '说明', '示例'],
    [
        ['feat', '新功能', 'feat: 实现 ContentAgent 小红书 SOP 流程'],
        ['fix', 'Bug 修复', 'fix: 修复 SSE 流式输出时 metadata 为空导致 NPE 的问题'],
        ['refactor', '代码重构', 'refactor: 将 LLMClient 抽象接口从 agent-core 提取到 llm 包'],
        ['docs', '文档更新', 'docs: 补充系统设计文档技术关键点章节'],
        ['style', '代码风格调整', 'style: 统一 .vue 文件缩进为 2 空格'],
        ['test', '测试', 'test: 添加 AgentRouter 关键词路由单元测试'],
        ['chore', '构建/工具/配置', 'chore: 更新 .gitignore 增加 Python 虚拟环境忽略规则'],
    ]
)

add_heading('3.2  目录与模块约定', 2)

add_para('项目采用前后端分离的目录结构，各模块职责明确，避免代码混杂：')

add_para('后端模块结构（backend/）：', bold=True)
add_table(
    ['模块目录', '对应 Maven artifact', '职责'],
    [
        ['app-service/', 'app-service', '聚合启动入口，包含 FlowMindApplication 主类'],
        ['gateway-service/', 'gateway-service', '统一路由、CORS 配置'],
        ['user-service/', 'user-service', '用户认证、角色权限管理'],
        ['ai-agent-service/', 'ai-agent-service', 'AI 智能体核心：Agent Router、LLMClient、Extension'],
        ['content-service/', 'content-service', '内容运营：Port/Adapter 架构、SOP 引擎'],
        ['knowledge-service/', 'knowledge-service', '知识库管理、向量检索、Weaviate 集成'],
        ['student-service/', 'student-service', '学员画像与申请进度'],
        ['school-service/', 'school-service', '院校信息与项目情报'],
        ['analytics-service/', 'analytics-service', '数据分析与图表接口'],
        ['feishu-service/', 'feishu-service', '飞书能力封装与 Mock'],
        ['common-core/', 'common-core', '公共工具：ApiResponse、IdGenerator'],
        ['common-security/', 'common-security', '安全基础：AuthInterceptor、TokenUtil'],
        ['common-web/', 'common-web', 'Web 配置：OpenApiConfig、CORS'],
    ]
)

add_para('前端目录结构（frontend/src/）：', bold=True)
add_table(
    ['目录', '职责'],
    [
        ['api/', 'API 客户端封装（agent.ts / client.ts / content.ts / knowledge.ts / knowledge-sync.ts / mock.ts / safe-request.ts，共 7 个模块）'],
        ['components/', '可复用 Vue 组件（AgentChat / ChatMessage / SessionList / AgentSidebar / ContextPanel / ContentCard / ContentSopPanel / StarRating / StatCard / ChartCard / SchoolProjectCard / StudentTable / PermissionGate / SafetyDiagnosticsPanel / SafeStateView / CapabilityParityPanel / OperationChecklist / ResponsiveDataPanel / ConversationHistory / AgentTabs，共 20 个）'],
        ['layouts/', '布局组件（MainLayout.vue）'],
        ['router/', 'Vue Router 路由配置与守卫（含权限判断 beforeEach）'],
        ['stores/', 'Pinia 状态管理（auth.ts，含用户信息、Token、角色权限）'],
        ['utils/', '工具库（datetime.ts / feature-parity.ts / form-safety.ts / guardrails.ts / markdown.ts / page-safety-blueprints.ts，共 6 个安全与功能工具）'],
        ['views/', '页面视图组件（LoginView / DashboardView / AgentWorkspaceView / ContentView / KnowledgeView / StudentsView / SchoolsView / AnalyticsView / FeishuView / SettingsView，共 10 个业务页面）'],
        ['assets/', '静态资源'],
    ]
)

add_para('其他重要目录：', bold=True)
add_table(
    ['目录', '说明'],
    [
        ['app/', 'Android 移动客户端（Java 原生，含 MainActivity + ApiClient + Gradle 构建系统）'],
        ['desktop_fronted/', '桌面客户端（Python/PySide6，高保真复刻 Web 端全部 10 个页面，含离线 Mock 模式）'],
        ['docs/', '项目文档：需求文档、设计文档、API 文档、架构图、提示词等'],
        ['docs/submission/', '课程提交文档：系统需求文档、系统设计文档、项目管理文档'],
        ['backend/sql/', '数据库初始化脚本（schema.sql 19 张表 + mock-data.sql 示例数据）'],
        ['scripts/', '辅助脚本（含系统设计文档和项目管理文档的自动生成脚本）'],
        ['.gitignore', 'Git 忽略规则（含 Python .venv、__pycache__、Node modules、Maven target、application-local.yml 等）'],
    ]
)

add_heading('3.3  扩展能力开发规范', 2)
add_para('为了降低多人协作时的代码冲突，项目设计了 contrib/capability 扩展缓冲区机制。'
         '当前仓库中已落地的扩展能力包括：')
add_para('① VectorSearchRuntimeExtension：位于 backend/app-service/src/main/java/com/flowmind/contrib/capability/vectorsearch/，'
         '实现 AgentExtension 接口，提供向量检索的运行时上下文注入。该扩展在 ai-agent-service 的 DefaultAgentExtensions 中被自动发现和注册。')
add_para('② 新增扩展能力的标准流程：在 contrib/capability 包下创建独立的能力类 → 实现 AgentExtension 接口（name / type / description / supports / runtimeContext 五个方法）'
         '→ 注册为 Spring @Component → Agent Router 和 DefaultAgentExtensions 自动发现并加载。')
add_para('③ 能力稳定并通过测试后，由组长将其迁移到正式的 ai-agent-service/src/main/java/com/flowmind/agent/extension/ 包。'
         '当前该包下已包含：AgentExtension（接口）、McpToolProvider、SkillProvider、LarkCliMcpExtension、'
         'SemanticVectorSearchExtension、RuntimeToolExtensions、DefaultAgentExtensions 共 7 个扩展组件。')
add_para('④ 类似地，桌面客户端（desktop_fronted/）和 Android 移动客户端（app/）虽然不在 contrib 包中，'
         '但同样遵循"独立目录开发、通过统一 API 与后端交互、不修改后端主流程"的隔离原则，'
         '使得前端、桌面、移动三端的开发可以完全并行，互不阻塞。')
add_para('⑤ 这种多层次隔离机制确保：主流程代码（AgentRouter、BaseAgent、LLMClient）不被频繁修改，'
         '降低合并冲突概率；每个成员可以独立开发、测试自己的扩展能力；'
         '新增一个 MCP 工具或 Skill 不需要修改任何已有 Agent 的代码。')

add_heading('3.4  沟通与协作工具', 2)
add_table(
    ['工具', '用途', '说明'],
    [
        ['GitHub', '代码管理 + Issue 追踪', '所有代码变更通过 GitHub 管理，建议利用 Issues 追踪任务和 Bug'],
        ['飞书 / 微信群', '日常沟通与同步', '紧急问题和每日进度同步'],
        ['飞书文档', '会议记录与设计讨论', '重要决策和技术方案讨论通过飞书文档沉淀'],
        ['线下会议', '阶段回顾与评审', '需求分析、设计评审、集成测试等关键节点进行线下讨论'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 第四章 开发计划与进度管理 ═══════════════════════════
add_heading('4  开发计划与进度管理', 1)

add_heading('4.1  项目整体里程碑', 2)
add_para('项目周期为 2026年5月20日至6月20日，共31天。分为6个阶段：')

add_table(
    ['阶段', '时间', '核心任务', '里程碑产出物'],
    [
        ['M1 需求分析', '5/20 - 5/24\n（5天）', '领域边界明确、用户角色识别、内容 SOP 梳理、竞品分析、用例建模', '需求文档初稿、用例图初稿'],
        ['M2 软件设计', '5/25 - 5/31\n（7天）', '系统架构设计、前后端职责划分、数据库模型设计、Agent 抽象设计、接口清单', '设计文档初稿、包图/类图、数据库设计'],
        ['M3 基础开发', '6/1 - 6/7\n（7天）', '前后端框架搭建、登录与主布局、Mock 数据准备、基础 README', '可运行的骨架系统'],
        ['M4 核心开发', '6/8 - 6/14\n（7天）', 'AI 工作台、内容运营、知识库+向量检索、飞书同步、学员/院校/数据分析', '核心功能全部完成'],
        ['M5 集成测试', '6/15 - 6/18\n（4天）', '前后端联调、流式回复/Thinking/Trace 优化、安全配置检查、UI 响应式优化', '稳定可演示版本'],
        ['M6 答辩准备', '6/19 - 6/20\n（2天）', '运行截图、演示脚本、Git 仓库检查、PPT 整理、分工彩排', '完整提交物'],
    ]
)

add_heading('4.2  详细任务分解（WBS）', 2)

add_para('M1 需求分析阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['1.1', '明确系统领域边界与范围', '待补充', '1天', '—'],
        ['1.2', '识别用户角色与需求', '待补充', '1天', '1.1'],
        ['1.3', '梳理内容运营 SOP（小红书/朋友圈/知识库）', '待补充', '1天', '1.1'],
        ['1.4', '梳理知识库、飞书、学员、院校、分析模块需求', '待补充', '1天', '1.2'],
        ['1.5', '竞品分析', '待补充', '0.5天', '1.2'],
        ['1.6', '撰写需求文档初稿与用例图', '待补充', '1天', '1.3, 1.4, 1.5'],
    ]
)

add_para('M2 软件设计阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['2.1', '系统总体架构设计', '待补充', '2天', '1.6'],
        ['2.2', '前后端职责与接口划分', '待补充', '1天', '2.1'],
        ['2.3', '数据库模型设计（MySQL + Weaviate）', '待补充', '2天', '2.1'],
        ['2.4', 'Agent 抽象、LLMClient、Extension 扩展设计', '待补充', '2天', '2.1'],
        ['2.5', '包图、领域类图、接口清单输出', '待补充', '2天', '2.2, 2.3, 2.4'],
        ['2.6', '设计文档初稿撰写', '待补充', '2天', '2.5'],
    ]
)

add_para('M3 基础开发阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['3.1', '搭建后端 Maven 多模块骨架', '待补充', '1天', '2.6'],
        ['3.2', '搭建前端 Vue 3 项目骨架', '待补充', '1天', '2.6'],
        ['3.3', '编写数据库初始化脚本（schema.sql + mock-data.sql）', '待补充', '1天', '3.1'],
        ['3.4', '实现登录/注册/主布局/路由守卫', '待补充', '2天', '3.1, 3.2'],
        ['3.5', '实现 gateway-service 统一路由与 CORS', '待补充', '0.5天', '3.1'],
        ['3.6', '准备 Swagger/OpenAPI 文档配置', '待补充', '0.5天', '3.1'],
        ['3.7', '编写基础 README.md', '待补充', '0.5天', '3.3'],
    ]
)

add_para('M4 核心开发阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['4.1', '实现 AI 工作台：Agent Router、LLMClient、SSE 流式输出', '待补充', '3天', '3.4'],
        ['4.2', '实现 5 个 Agent（Content/Knowledge/Student/School/Feishu）', '待补充', '3天', '4.1'],
        ['4.3', '实现 AgentExtension 扩展机制（MCP/Skill/Runtime）', '待补充', '2天', '4.1'],
        ['4.4', '实现内容运营模块：Port/Adapter + SOP + 主题库/文案库/日历', '待补充', '3天', '3.4'],
        ['4.5', '实现知识库模块：Weaviate 集成 + Embedding + 向量检索', '待补充', '3天', '3.4'],
        ['4.6', '实现飞书同步模块：lark-cli 桥接 + Mock + 同步日志', '待补充', '2天', '3.4'],
        ['4.7', '实现学员管理 + 院校情报 + 数据分析模块', '待补充', '3天', '3.4'],
        ['4.8', '前端业务页面开发（内容/知识库/学员/院校/数据分析/飞书/设置）', '待补充', '4天', '3.4'],
        ['4.9', '前端 AI 工作台 SSE 流式展示 + Thinking/Trace 面板', '待补充', '2天', '4.1'],
        ['4.10', '前端 ECharts 数据分析图表', '待补充', '1天', '4.8'],
        ['4.11', 'desktop_fronted 桌面客户端开发', '待补充', '3天', '3.4'],
    ]
)

add_para('M5 集成测试阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['5.1', '前后端联调（API 接口对齐、错误处理）', '待补充', '1.5天', 'M4 全部'],
        ['5.2', 'SSE 流式回复稳定性测试', '待补充', '0.5天', '4.1, 4.9'],
        ['5.3', 'Thinking 与 Trace 展示优化', '待补充', '0.5天', '4.1, 4.9'],
        ['5.4', '数据库初始化与 Mock 数据检查', '待补充', '0.5天', '3.3'],
        ['5.5', '安全配置检查（API Key 不泄露、.gitignore 完整性）', '待补充', '0.5天', '—'],
        ['5.6', 'UI 响应式与浏览器兼容性检查', '待补充', '0.5天', 'M4 全部'],
        ['5.7', '最终 Bug 修复', '待补充', '1天', '5.1-5.6'],
    ]
)

add_para('M6 答辩准备阶段：', bold=True)
add_table(
    ['任务ID', '任务名称', '负责人', '工作量', '前置任务'],
    [
        ['6.1', '截取各页面运行截图', '待补充', '0.5天', '5.7'],
        ['6.2', '撰写系统需求文档终稿', '待补充', '1天', '5.7'],
        ['6.3', '撰写系统设计文档终稿', '待补充', '1天', '5.7'],
        ['6.4', '撰写项目管理文档终稿', '待补充', '0.5天', '5.7'],
        ['6.5', '准备演示脚本与 PPT', '待补充', '1天', '6.2, 6.3'],
        ['6.6', 'Git 仓库最终检查与清理', '待补充', '0.5天', '6.4'],
        ['6.7', '答辩彩排', '待补充', '0.5天', '6.5'],
    ]
)

add_heading('4.3  进度跟踪机制', 2)
add_para('项目进度跟踪采用以下方式：')
add_para('① 每日简短同步：每名成员在开始开发前简要同步昨日进展和当日计划，遇阻塞问题及时沟通。')
add_para('② GitHub Issues：每个功能模块创建对应的 Issue，标注负责人和预计完成时间，完成后关闭。')
add_para('③ 里程碑检查：每个阶段结束时，组长组织一次简短检查，确认该阶段产出物是否达标。')
add_para('④ 风险及时上报：如某个任务因技术困难或依赖阻塞无法按时完成，成员应立即上报，由组长协调资源或调整计划。')

doc.add_page_break()

# ═══════════════════════════ 第五章 风险管理 ═══════════════════════════
add_heading('5  风险管理', 1)

add_heading('5.1  风险识别与应对策略', 2)

add_table(
    ['风险编号', '风险描述', '可能性', '影响程度', '应对策略'],
    [
        ['R1', 'LLM API Key 不可用或余额不足，导致 AI 功能无法演示',
         '中', '高', '已实现 MockLLMClient 作为降级方案，无 Key 时自动切换为 Mock 模式，返回预设的模拟数据'],
        ['R2', '飞书 API 权限不足，无法调用真实飞书接口',
         '中', '中', '已实现 FeishuMockClient 和 LarkCliMcpExtension 的 Mock 模式，无飞书权限时可演示飞书操作的模拟流程'],
        ['R3', 'Weaviate Docker 容器启动失败或向量检索异常',
         '低', '中', '向量检索失败时可降级为 MySQL LIKE 文本搜索；系统启动时自动检查 Weaviate 连接状态'],
        ['R4', '前端 SSE 流式展示出现兼容性问题（某些浏览器不支持 EventSource）',
         '低', '低', 'Chrome/Edge 均原生支持 SSE；若出现兼容性问题，可降级为非流式的"等待-完整返回"模式'],
        ['R5', '多人协作时 Git 合并冲突复杂，影响开发进度',
         '中', '中', '通过功能分支隔离开发；contrib/capability 缓冲机制减少主流程修改；组长负责解决冲突'],
        ['R6', '成员对技术栈（Vue 3 / Spring Boot / Weaviate）不熟悉，学习成本高',
         '中', '中', '需求分析阶段预留学习缓冲；模块按难度分配；提供代码模板和示例'],
        ['R7', '课程设计时间紧张，部分功能无法完成',
         '高', '高', '按优先级开发，核心功能（AI 工作台 + 内容运营 + 知识库）优先；非核心功能可标注"后续扩展"'],
        ['R8', '同学在答辩时网络环境不佳，无法连接 LLM API 或 Weaviate',
         '中', '高', 'MockLLMClient 支持完全离线 Demo；确保所有核心流程可在纯 Mock 模式下完整演示'],
    ]
)

add_heading('5.2  风险监控与报告', 2)
add_para('① 每个开发阶段初期进行风险回顾，检查之前识别的风险是否已解除或变化。')
add_para('② 新风险出现时立即记录，评估影响并制定应对方案。')
add_para('③ 高风险项（如 R7、R8）在答辩前进行专项复查，确保有可靠的降级方案。')

doc.add_page_break()

# ═══════════════════════════ 第六章 质量保证计划 ═══════════════════════════
add_heading('6  质量保证计划', 1)

add_heading('6.1  代码质量控制', 2)
add_para('① 代码风格统一：后端 Java 代码遵循 Spring Boot 通用规范（Controller-Service-Mapper 分层）；'
         '前端 Vue 代码遵循 Vue 3 Composition API 风格，使用 TypeScript 定义接口类型。')
add_para('② 代码审查：组长对合并到 main 分支的所有代码进行审查，重点关注：接口设计合理性、'
         '异常处理完整性、敏感信息不泄露、模块边界是否清晰。')
add_para('③ 敏感信息检查：每次合并前确认：application-local.yml 不进入 Git、'
         '前端代码中无硬编码的 API Key、飞书 Token 等敏感信息。')
add_para('④ 构建检查：提交前确保后端 Maven 编译通过（mvnw compile），前端 Vite 构建通过（npm run build）。')

add_heading('6.2  功能测试要点', 2)
add_para('核心功能测试清单：')

add_table(
    ['测试项', '测试内容', '通过标准'],
    [
        ['登录/注册', '正常登录、错误密码、未登录拦截、注册重复用户名', '登录成功获取 Token、错误有提示、未登录跳转登录页'],
        ['AI 工作台问答', '输入问题 → SSE 流式返回 → 显示回答', '流式输出正常、回答内容合理（或 Mock 数据合理）'],
        ['Agent 自动路由', '输入不同类型问题（内容/知识/学员/院校/飞书）', '自动识别意图并路由到对应 Agent'],
        ['Thinking 展示', '模型返回思考过程', '前端可折叠展示 Thinking 区域'],
        ['Trace 展示', '工具调用结果', '前端可展开查看工具调用的输入输出'],
        ['内容主题库', '新增/删除/评分/查看主题', 'CRUD 正常、评分正常保存、历史文案可查看'],
        ['知识库检索', '输入查询文本 → 向量检索 → 返回结果', '返回相关文档及相似度、无结果时正确降级'],
        ['学员管理', '新增/编辑/删除学员、申请进度管理', 'CRUD 正常、数据正确保存'],
        ['飞书同步', '飞书文档同步、同步日志', 'Mock 模式下返回模拟同步结果'],
        ['数据分析', '各图表数据加载与展示', 'ECharts 图表正常渲染、数据与数据库一致'],
        ['会话历史', '切换会话、查看历史消息', '历史消息完整展示、Thinking/Trace 可回看'],
    ]
)

add_heading('6.3  文档质量控制', 2)
add_para('① 需求文档、设计文档、项目管理文档三份核心文档需要相互一致（项目名称、模块划分、技术选型等）。')
add_para('② 文档中不出现真实 API Key、密码、隐私数据。')
add_para('③ 文档标注"待补充"的部分在最终提交时统一填写。')
add_para('④ 图表需要图名、图注和说明文字。')

doc.add_page_break()

# ═══════════════════════════ 第七章 配置管理 ═══════════════════════════
add_heading('7  配置管理', 1)

add_heading('7.1  版本控制策略', 2)
add_para('项目所有源代码、配置文件（除敏感配置外）、SQL 脚本、文档均纳入 Git 版本管理。')
add_para('使用 .gitignore 排除以下内容：')
add_table(
    ['排除项', '说明'],
    [
        ['node_modules/', '前端依赖包'],
        ['**/target/', 'Maven 构建产物'],
        ['.venv/', 'Python 虚拟环境'],
        ['__pycache__/', 'Python 字节码缓存'],
        ['*.pyc', 'Python 编译文件'],
        ['backend/.m2repo/', 'Maven 本地仓库'],
        ['dist/', '前端构建产物'],
        ['application-local.yml', '包含敏感信息的本地配置文件'],
        ['.env', '环境变量文件（如有）'],
    ]
)

add_heading('7.2  配置项管理', 2)
add_para('项目配置分为三层：')
add_para('① 公共默认配置（application.yml）：端口、数据源类型、Weaviate 地址等非敏感信息，进入 Git。')
add_para('② 本地私有配置（application-local.yml）：数据库密码、LLM API Key、飞书 App Secret，不进入 Git。'
         '提供 application-local.yml.example 模板文件，包含配置项名称和占位符。')
add_para('③ 前端环境变量（.env / .env.development）：VITE_API_BASE_URL 等前端配置，'
         '不包含任何 Key 或 Secret。')

add_heading('7.3  发布版本管理', 2)
add_para('Demo 阶段的版本标签：')
add_para('• v1.0.0 = FlowMindAgent1.0.0：课程设计提交版本，包含全部核心功能和文档。')
add_para('• 后续如有 Bug 修复，按 v1.0.1、v1.0.2 递增。')
add_para('• 如有重大功能更新，按 v1.1.0、v1.2.0 递增。')

doc.add_page_break()

# ═══════════════════════════ 第八章 交付物清单 ═══════════════════════════
add_heading('8  交付物清单', 1)

add_table(
    ['序号', '交付物', '格式', '说明', '状态'],
    [
        ['1', '系统需求文档', 'Word (.docx)', '包含项目可行性分析、功能分析、原型设计、开发计划、领域模型、动态模型、非功能性需求等', '待定稿'],
        ['2', '系统设计文档', 'Word (.docx)', '包含技术选型、系统架构、组件与模块设计、数据架构、服务协作、安全设计、部署方案、技术关键点等', '待定稿'],
        ['3', '项目管理文档', 'Word (.docx)', '包含团队分工、协作规范、开发计划、风险管理、质量保证、配置管理等', '待定稿'],
        ['4', '项目源代码', 'GitHub 仓库', '包含完整的前端、后端、桌面客户端源代码及配置文件', '已完成'],
        ['5', '数据库脚本', 'SQL 文件', 'backend/sql/schema.sql + mock-data.sql', '已完成'],
        ['6', 'README.md', 'Markdown', '项目说明、启动步骤、技术栈、目录结构', '已完成'],
        ['7', 'API 文档', 'Swagger/OpenAPI', '自动生成的接口文档', '已完成'],
        ['8', '运行截图', 'PNG/JPG', '各页面实际运行截图（登录、AI 工作台、内容管理、知识库、学员管理等）', '待补充'],
        ['9', '演示 PPT', 'PPTX', '答辩演示幻灯片', '待补充'],
        ['10', '演示脚本', '文本/Word', '答辩演示的流程脚本', '待补充'],
        ['11', 'Git 提交历史', 'Git Log', '可展示团队协作过程和开发节奏', '已完成'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 第九章 沟通记录 ═══════════════════════════
add_heading('9  关键决策与沟通记录', 1)

add_para('以下记录项目过程中的关键决策和讨论结论（待补充日期和细节）：')

add_table(
    ['序号', '日期', '议题', '结论/决策', '参与人'],
    [
        ['1', '5月下旬', '系统技术栈选型',
         '确定前端采用 Vue 3 + Element Plus + Vite，后端采用 Spring Boot 3 + Maven 多模块，数据存储采用 MySQL + Weaviate',
         '待补充'],
        ['2', '5月下旬', 'Agent 架构设计',
         '确定 Agent Router 自动路由 + LLMClient 抽象 + AgentExtension 扩展的三层架构，"不绑定大模型厂商、不解耦工具能力"的设计原则',
         '待补充'],
        ['3', '5月下旬', '内容运营模块架构',
         '确定 content-service 采用 Port/Adapter 模式，将内容生成、知识检索、飞书发布、数据存储定义为四个抽象端口，Mock 和 MySQL 作为当前适配器实现',
         '待补充'],
        ['4', '6月初', 'Demo 聚合启动 vs 真实微服务',
         'Demo 阶段使用 app-service 聚合启动以降低运行复杂度，但源码层面保留清晰的微服务边界，后续可平滑拆分',
         '待补充'],
        ['5', '6月上旬', '敏感信息管理',
         '确定使用 application-local.yml（已加入 .gitignore）保存敏感配置；前端不持有任何 LLM Key 或飞书 App Secret；提供 application-local.yml.example 模板',
         '待补充'],
        ['6', '6月上旬', 'Contrib/Capability 扩展缓冲区',
         '确定在 app-service 下设立 contrib/capability 包作为成员的独立开发缓冲区，稳定后迁移到正式模块',
         '待补充'],
        ['7', '6月中旬', 'SSE vs WebSocket',
         '确定使用 SSE 作为 AI 流式输出的协议，因其单向推送特征与 LLM 流式生成场景完美匹配，且实现更轻量',
         '待补充'],
        ['8', '6月中旬', '飞书接入方案',
         '确定当前使用 lark-cli 命令行桥接 + FeishuMockClient Mock 模式的方案，后续可替换为原生飞书 OpenAPI SDK',
         '待补充'],
    ]
)

doc.add_page_break()

# ═══════════════════════════ 第十章 总结与待补充 ═══════════════════════════
add_heading('10  总结与待补充项', 1)

add_heading('10.1  项目总结', 2)
add_para(
    'FlowMind Agent 项目在 31 天的课程设计周期内，从零开始构建了一个包含前端 AI 工作台、后端 12 个服务模块、'
    '向量检索引擎、飞书生态集成、桌面客户端等完整功能栈的 AI 智能体平台。'
    '项目在软件工程实践方面体现了以下特点：'
)
add_para('① 遵循了完整的软件工程生命周期：从需求分析、系统设计、编码实现到集成测试和文档交付。')
add_para('② 采用了模块化、接口抽象、Port/Adapter 等良好的软件架构设计模式，使系统具备了可维护性和可扩展性。')
add_para('③ 合理运用了 Git 分支管理、功能隔离、代码审查等团队协作实践。')
add_para('④ 输出了需求文档、设计文档、项目管理文档三份完整的工程文档。')
add_para('⑤ 在 Demo 阶段实现了 AI 工作台、Agent 自动路由、LLM API 可替换、工具扩展解耦、SSE 流式输出、'
         '向量检索、内容 SOP、飞书集成等多个技术关键点。')

add_heading('10.2  待补充项', 2)
add_para('以下内容需要在最终提交时补充完善：')

add_table(
    ['序号', '待补充内容', '说明', '负责人'],
    [
        ['1', '团队成员姓名与分工', '第 2 章中所有"待补充"的人员信息', '待补充'],
        ['2', '工作量统计', '各模块的实际工作量（人天）', '待补充'],
        ['3', '代码量统计', '使用 cloc 工具统计前端、后端、桌面客户端、文档的代码行数', '待补充'],
        ['4', '各页面运行截图', '登录页、AI 工作台、内容管理、知识库、学员管理、院校情报、数据分析、飞书同步、系统设置', '待补充'],
        ['5', '功能测试结果', '第 6 章功能测试清单的实际执行结果（通过/未通过）', '待补充'],
        ['6', '沟通记录日期与参与人', '第 9 章关键决策记录的具体日期和参与人', '待补充'],
        ['7', '答辩 PPT', '准备答辩演示用的 PPT 文件', '待补充'],
        ['8', '演示脚本', '准备答辩演示流程脚本', '待补充'],
        ['9', '最终代码检查', '确认 .gitignore 完整、无敏感信息泄露、所有"待补充"已填写', '待补充'],
    ]
)

add_heading('10.3  后续优化方向', 2)
add_para('课程设计提交后，系统可在以下方向继续优化（非本次提交要求）：')
add_para('① 将 app-service 聚合启动拆分为真正的微服务部署（Spring Cloud Gateway + Nacos）。')
add_para('② 引入 Redis 缓存和 MinIO 对象存储，提升性能和存储能力。')
add_para('③ 将 lark-cli 桥接替换为原生飞书 OpenAPI Java SDK，提升稳定性和性能。')
add_para('④ 升级 Token 认证为 JWT + RBAC 完整权限体系。')
add_para('⑤ 引入 Prometheus + Grafana 监控体系和 CI/CD 自动化部署。')
add_para('⑥ 优化 Agent Router 的关键词匹配为基于 LLM 的意图分类。')
add_para('⑦ 支持内容 SOP 的自定义配置和运营人员的 Prompt 模板编辑。')

doc.add_page_break()

# ═══════════════════════════ 保存 ═══════════════════════════
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'docs', 'submission', '项目管理文档.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
